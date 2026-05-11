from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

def set_para_format(para, alignment=None, bold=False, italic=False, font_size=12, font_name='Times New Roman', space_after=Pt(6)):
    para_format = para.paragraph_format
    if alignment:
        para.alignment = alignment
    para_format.space_after = space_after
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))
        return p
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(10))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(8))
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(6))
        return p

def add_normal(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Inches(0.5), bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_para_format(p, alignment=alignment, space_after=Pt(6))
    p.paragraph_format.first_line_indent = first_line_indent
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
    p.paragraph_format.first_line_indent = Inches(0)
    return p

def add_table_grid(rows_data, header_row):
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    for i, txt in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows_data:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            row.cells[i].text = txt
    return table

# ===== TITLE PAGE =====
p = doc.add_paragraph()
run = p.add_run('V-BUS — REAL-TIME SMART BUS TRACKING AND PASSENGER INFORMATION SYSTEM')
run.bold = True
run.font.size = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))

add_normal('A PROJECT REPORT', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))
add_normal('Submitted by', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))
add_normal('in partial fulfilment for the award of the degree', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))
add_normal('of', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))
p = doc.add_paragraph()
run = p.add_run('BACHELOR OF ENGINEERING')
run.bold = True
run.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(8))
add_normal('in', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))
p = doc.add_paragraph()
run = p.add_run('COMPUTER SCIENCE AND ENGINEERING')
run.bold = True
run.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(8))

add_normal('THANTHAI PERIYAR GOVERNMENT INSTITUTE OF TECHNOLOGY VELLORE - 632002', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))
add_normal('ANNA UNIVERSITY : CHENNAI 600 025', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))
add_normal('MAY 2026', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Inches(0))

doc.add_page_break()

# ===== BONAFIDE CERTIFICATE =====
add_heading_custom('BONAFIDE CERTIFICATE', level=1)
add_normal('Certified that this project report “V-BUS — REAL-TIME SMART BUS TRACKING AND PASSENGER INFORMATION SYSTEM” is the bonafide work of “VIJAY VASANTH A K S (513122104051), AKASHH D (513122104001), KAVIN E (513121104019)”, who carried out the project work under my supervision.')
add_normal('')
add_normal('')
add_normal('')
add_normal('Submitted for the Project Viva-Voce Examination held on ______________', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('')
add_normal('')
add_normal('INTERNAL EXAMINER                                    EXTERNAL EXAMINER', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

doc.add_page_break()

# ===== ACKNOWLEDGEMENT =====
add_heading_custom('ACKNOWLEDGEMENT', level=1)
add_normal('The successful completion of this project would not have been possible without the support, guidance, and encouragement of many individuals, to whom we express our sincere gratitude.')
add_normal('We extend our heartfelt thanks to Dr. P. K. Palani, B.E. (Hons), M.E., Ph.D., Principal of Thanthai Periyar Government Institute of Technology, Vellore, for his valuable motivation and constant support throughout our academic journey.')
add_normal('We are deeply grateful to Dr. S. Letitia, M.E., Ph.D., Head of the Department of Computer Science and Engineering, for her insightful suggestions and unwavering encouragement that greatly enriched our work.')
add_normal('I hereby express my deep sense of gratitude to the project coordinator Dr. N. Jagadeeswari, M.E., Ph.D., Assistant Professor of the Department of Computer Science and Engineering, Thanthai Periyar Government Institute of Technology, Vellore, for expert guidance and encouragement throughout the project.')
add_normal('Our sincere appreciation goes to our project guide, Mrs. N. Naveenabegum, M.E., Assistant Professor [Adhoc], Department of Computer Science and Engineering, for her expert guidance, continuous support, and valuable feedback that helped us complete this project successfully.')
add_normal('We would also like to thank all the faculty members of the Department of Computer Science and Engineering for their constant support and encouragement, which played a vital role in the completion of our project.')

doc.add_page_break()

# ===== ABSTRACT =====
add_heading_custom('ABSTRACT', level=1)
add_normal('The rapid growth of urban transportation systems has created a pressing need for intelligent, real-time monitoring solutions that improve passenger convenience and operational efficiency. Traditional public transport systems often lack accurate information about bus locations, arrival times, and route conditions, leading to uncertainty, prolonged waiting periods, and reduced commuter satisfaction. To address these challenges, this project proposes V-Bus, a Real-Time Smart Bus Tracking and Passenger Information System that leverages modern web technologies, event-driven architecture, and real-time communication to provide seamless tracking and management of buses.')
add_normal('The system is designed to capture live location data from bus drivers through a dedicated mobile application built with React Native and Expo. Location updates are transmitted instantly to a centralized backend server using Socket.IO, a low-latency, bidirectional communication protocol. The backend, developed using Node.js and Express.js, maintains an authoritative tracking state, processes incoming location streams, and synchronizes updates across all connected passenger applications in real time. Data persistence is handled by MongoDB, a scalable NoSQL database that stores route geometries, stop configurations, bus assignments, and historical trip records.')
add_normal('The passenger application provides an immersive tracking experience through an embedded WebView rendering interactive Leaflet.js maps over OpenStreetMap tiles. Passengers can view full map and mini map visualizations, observe live bus markers, follow selected buses, monitor route corridors, and receive accurate estimated time of arrival (ETA) calculations powered by a route-aware projection engine. The system includes advanced features such as nearby stop detection, real-time popup updates, automatic offline cleanup, and a follow-bus mode that keeps the viewport centered on the selected vehicle.')
add_normal('For drivers, the application offers secure login, route assignment, trip lifecycle management, background GPS tracking, and an SOS emergency button for immediate alerting. The backend enforces rigorous state management through a tracking state map, handles BUS_OFFLINE events to remove stale data, and delivers consistent, synchronized updates to all clients through a single socket architecture.')
add_normal('By integrating a production-grade backend with scalable real-time communication and a responsive cross-platform mobile interface, V-Bus ensures high performance, reliability, and accessibility. It reduces passenger waiting time, enhances transport management visibility, and contributes to the development of smarter urban mobility solutions. Future enhancements may include AI-based traffic prediction, multi-city deployment, and deeper integration with smart city infrastructure.')

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
add_heading_custom('TABLE OF CONTENTS', level=1)
add_normal('1\tINTRODUCTION\t1', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t1.1 Problem Statement\t2', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t1.2 Objectives\t2', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t1.3 Scope\t3', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('2\tLITERATURE SURVEY\t4', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t2.1 GPS-Based Tracking Systems\t4', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t2.2 Real-Time Communication in Transportation\t4', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t2.3 Event-Driven Vehicle Monitoring Systems\t5', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t2.4 Mobile-Based Tracking Applications\t5', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t2.5 WebView and Hybrid Map Rendering\t6', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t2.6 Survey on Smart Transport Technologies\t6', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('3\tEXISTING AND PROPOSED SYSTEM\t7', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t3.1 Existing System\t7', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.1.1 Manual Bus Tracking\t7', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.1.2 GPS Tracking without Integration\t7', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.1.3 Basic Mobile Tracking Apps\t7', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.1.4 Limitations of Existing System\t8', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t3.2 Proposed System\t9', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.2.1 System Overview\t9', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.2.2 Core Modules and Technologies\t9', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.2.3 Key Features\t10', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t3.2.4 Advantages Over Existing Systems\t11', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('4\tSYSTEM SPECIFICATION\t12', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t4.1 Hardware Requirements\t12', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t4.2 Software Requirements\t13', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t4.3 Network Requirements\t14', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t4.4 Security Requirements\t14', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('5\tSYSTEM DESIGN\t15', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t5.1 UML Diagrams\t15', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t5.1.1 Use Case Diagram\t15', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t5.1.2 Activity Diagram\t16', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t5.1.3 Sequence Diagram\t17', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t5.1.4 Class Diagram\t18', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t5.1.5 Architecture Diagram\t19', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('6\tSYSTEM MODULES\t21', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t6.1 Real-Time Tracking Module\t22', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t6.2 Driver Application Module\t22', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t6.3 Backend API and Socket.IO Module\t23', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t6.4 Passenger Application Module\t24', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t6.5 Route-Aware Projection Engine Module\t25', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t6.6 Database and State Management Module\t25', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('7\tSYSTEM IMPLEMENTATION\t27', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t7.1 Project Structure\t27', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t7.2 Development and Supporting Tools\t28', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('8\tTESTING\t31', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t8.1 Testing Objectives\t31', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t8.2 Testing Strategies\t31', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t8.2.1 Unit Testing\t31', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t8.2.2 Integration Testing\t32', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t8.2.3 Functional Testing\t32', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t8.2.4 Performance Testing\t33', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t\t8.2.5 Usability Testing\t33', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\t8.3 Test Cases\t33', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('9\tRESULTS AND PERFORMANCE\t35', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('10\tCONCLUSION AND FUTURE ENHANCEMENTS\t36', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\tAPPENDIX 1: SAMPLE CODE\t38', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\tAPPENDIX 2: SCREENSHOTS\t44', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('\tREFERENCES\t46', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

doc.add_page_break()

# ===== LIST OF FIGURES =====
add_heading_custom('LIST OF FIGURES', level=1)
add_normal('FIGURE NO.\tFIGURE NAME\tPAGE NO.', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('5.1\tUse Case Diagram\t15', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('5.2\tActivity Diagram\t16', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('5.3\tSequence Diagram\t17', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('5.4\tClass Diagram\t18', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('5.5\tArchitecture Diagram\t19', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('6.1\tSystem Modules\t21', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('7.1\tProject Directory Structure\t27', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('7.2\tWeb Admin Dashboard Interface\t28', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('7.3\tMobile Application Interface\t29', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('7.4\tReal-Time Bus Tracking Map View\t30', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

doc.add_page_break()

# ===== LIST OF TABLES =====
add_heading_custom('LIST OF TABLES', level=1)
add_normal('TABLE NO.\tTABLE NAME\tPAGE NO.', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('3.1\tAdvantages Over Existing Systems\t11', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('4.1\tHardware Requirements\t12', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('4.2\tSoftware Requirements\t13', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('4.3\tSecurity Requirements\t14', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('4.4\tNetwork Requirements\t14', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('7.1\tDevelopment Tools\t29', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('7.2\tSupporting Tools and Libraries\t30', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('8.1\tFunctional Testing\t32', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))
add_normal('8.2\tTest Cases\t33', alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

doc.add_page_break()

# ===== LIST OF ABBREVIATIONS =====
add_heading_custom('LIST OF ABBREVIATIONS', level=1)
abbrs = [
    ('GPS', 'Global Positioning System'),
    ('IoT', 'Internet of Things'),
    ('API', 'Application Programming Interface'),
    ('UI', 'User Interface'),
    ('UX', 'User Experience'),
    ('JWT', 'JSON Web Token'),
    ('HTTP', 'HyperText Transfer Protocol'),
    ('HTTPS', 'HyperText Transfer Protocol Secure'),
    ('DB', 'Database'),
    ('NoSQL', 'Non-Relational Database'),
    ('CPU', 'Central Processing Unit'),
    ('RAM', 'Random Access Memory'),
    ('ETA', 'Estimated Time of Arrival'),
    ('OS', 'Operating System'),
    ('URL', 'Uniform Resource Locator'),
    ('JSON', 'JavaScript Object Notation'),
    ('SSL', 'Secure Sockets Layer'),
    ('TLS', 'Transport Layer Security'),
    ('Socket.IO', 'Real-Time Event-Based Communication Library'),
    ('REST', 'Representational State Transfer'),
    ('ODM', 'Object Data Modeling'),
    ('SOS', 'Emergency Alert Signal'),
    ('HTML', 'HyperText Markup Language'),
    ('CSS', 'Cascading Style Sheets'),
    ('DOM', 'Document Object Model'),
    ('WebView', 'Embedded Web Browser Component'),
    ('SPA', 'Single Page Application'),
    ('OSM', 'OpenStreetMap'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Abbreviation'
t.rows[0].cells[1].text = 'Full Form'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
for abbr, full in abbrs:
    row = t.add_row()
    row.cells[0].text = abbr
    row.cells[1].text = full

doc.add_page_break()

# ===== CHAPTER 1: INTRODUCTION =====
add_heading_custom('CHAPTER 1', level=1)
add_heading_custom('INTRODUCTION', level=2)
add_normal('Effective public transportation is a crucial component of modern urban infrastructure, enabling mobility, reducing traffic congestion, and supporting economic development. However, many existing public transport systems lack real-time monitoring and efficient communication mechanisms, making it difficult for passengers to obtain accurate information about bus locations and arrival times. This often leads to increased waiting time, uncertainty, and inconvenience for daily commuters.')
add_normal('With the rapid advancement of technologies such as the Global Positioning System (GPS), real-time communication protocols, and cross-platform mobile development frameworks, it has become possible to develop intelligent transportation systems that provide real-time tracking and monitoring capabilities. These technologies enable continuous collection and transmission of location data, allowing users to access live updates about bus movements through modern mobile applications. Such systems significantly enhance user experience and improve operational efficiency in public transportation.')
add_normal('V-Bus is designed to address these challenges by providing a reliable, scalable, and production-grade solution for real-time bus tracking and passenger information. The system integrates a React Native driver application with a Node.js and Express.js backend, utilizing Socket.IO for instantaneous bidirectional communication. Live location data is captured from the driver device, transmitted to the backend, and synchronized across all connected passenger applications through a single socket architecture. The passenger application, built with React Native and an embedded WebView, renders interactive Leaflet.js maps over OpenStreetMap tiles to provide an immersive and responsive tracking experience.')
add_normal('In addition to passenger benefits, the system assists transport authorities in managing routes, monitoring driver activity, and analyzing historical data for better decision-making. The backend maintains an authoritative tracking state, performs route-aware corridor projection, manages stop progression logic, and automatically cleans up offline buses to ensure data consistency. By bridging the gap between modern software architecture and public transportation, V-Bus aims to create a smarter, more efficient, and user-friendly public transport experience.')

add_heading_custom('1.1 PROBLEM STATEMENT', level=3)
add_normal('Public transportation systems in many cities lack efficient real-time tracking and communication mechanisms, making it difficult for passengers to obtain accurate information about bus locations and arrival times. Commuters often depend on fixed schedules or manual inquiries, which are unreliable due to traffic conditions, delays, and operational uncertainties. This results in increased waiting time, poor travel planning, and overall inconvenience for daily users of public transport.')
add_normal('Although some systems provide basic GPS tracking, they are often limited in functionality, lacking integration with user-friendly applications and real-time data accessibility. Many existing solutions do not offer features such as accurate estimated time of arrival (ETA), route corridor visualization, real-time notifications, or follow-bus mode, which are essential for enhancing user experience. Additionally, transport authorities face challenges in monitoring bus operations, managing routes efficiently, and responding to real-time issues due to the absence of a centralized, event-driven, and intelligent system.')
add_normal('Therefore, there is a need for a comprehensive and scalable Smart Bus Tracking System that can provide accurate real-time location updates, improve communication between passengers and transport systems, and enhance operational efficiency through advanced technologies such as Node.js, Socket.IO, React Native, and OpenStreetMap.')

add_heading_custom('1.2 OBJECTIVES', level=3)
add_normal('The primary objectives of this work are:')
add_bullet('To develop a real-time bus tracking system that accurately captures and updates the location of buses using GPS technology and instant socket communication.')
add_bullet('To design and implement an authoritative backend system using Node.js and Express.js that processes location data, manages tracking state, and provides secure access through RESTful APIs and Socket.IO events.')
add_bullet('To create a user-friendly mobile application using React Native and Expo that allows passengers to track buses, view routes, check estimated time of arrival (ETA), and interact with live map visualizations through an embedded WebView.')
add_bullet('To enable transport authorities to manage buses, routes, and drivers efficiently through an integrated admin panel and driver application.')
add_bullet('To ensure system reliability, scalability, and real-time performance across multiple platforms using event-driven architecture, route-aware projection, and automated offline cleanup mechanisms.')

add_heading_custom('1.3 SCOPE', level=3)
add_normal('This project focuses on developing V-Bus, a Real-Time Smart Bus Tracking and Passenger Information System that provides real-time tracking of buses using GPS and Socket.IO technologies. The system is designed to enable passengers to monitor the live location of buses, view route details, estimate arrival times, and interact with immersive map interfaces through mobile applications. It aims to improve the overall efficiency and reliability of public transportation systems by offering accurate and timely information.')
add_normal('The scope of the project also includes the development of a backend system that manages data related to buses, routes, drivers, and real-time tracking state. The system ensures secure communication using RESTful APIs and Socket.IO events, and supports features such as user authentication, real-time updates, data storage, and automatic cleanup. Additionally, an admin interface and driver application are provided to transport authorities for efficient monitoring and management of operations.')
add_normal('The current implementation primarily focuses on urban transportation systems and supports functionalities such as live tracking, ETA calculation, route rendering, nearby stop detection, follow-bus mode, trip lifecycle management, and SOS emergency alerting. However, the system is designed with scalability in mind, allowing future enhancements and integration with advanced technologies.')
add_normal('The extended scope of the system may include:')
add_bullet('Integration with AI for traffic prediction and route optimization.')
add_bullet('Support for multiple cities and large-scale deployment.')
add_bullet('Push notification systems for delays, arrivals, and emergencies.')
add_bullet('Integration with smart city infrastructure and public transport networks.')
add_bullet('Advanced analytics for performance monitoring and decision-making.')

doc.add_page_break()

print("Completed through Chapter 1")
doc.save(r'w:\Final year project\V-Bus_Report.docx')
