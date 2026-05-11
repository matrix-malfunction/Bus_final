from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document(r'w:\Final year project\V-Bus_Report.docx')

def set_para_format(para, alignment=None, bold=False, italic=False, font_size=12, font_name='Times New Roman', space_after=Pt(6)):
    para_format = para.paragraph_format
    if alignment:
        para.alignment = alignment
    para_format.space_after = space_after
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))
        return p
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(10))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(8))
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(6))
        return p

def add_normal(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Inches(0.5), bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_para_format(p, alignment=alignment, space_after=Pt(6))
    p.paragraph_format.first_line_indent = first_line_indent
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
    p.paragraph_format.first_line_indent = Inches(0)
    return p

def add_table_grid(rows_data, header_row):
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    for i, txt in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows_data:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            row.cells[i].text = txt
    return table

# ===== CHAPTER 6: SYSTEM MODULES =====
add_heading_custom('CHAPTER 6', level=1)
add_heading_custom('SYSTEM MODULES', level=2)
add_normal('V-Bus is composed of multiple interconnected modules that work together to provide real-time tracking and efficient transport management. Each module performs a specific function and contributes to the overall system performance. These modules ensure smooth data collection, processing, and user interaction across the platform.')
add_normal('Figure – 6.1: System Modules')

add_heading_custom('6.1 Real-Time Tracking Module', level=3)
add_normal('Objective: To provide accurate real-time tracking of buses using GPS and Socket.IO.')
add_normal('This module is responsible for tracking the live location of buses using GPS technology. It collects location data such as latitude and longitude from the driver mobile application and updates it continuously through Socket.IO events. The data is transmitted to the backend server, where it is processed by the route-aware projection engine and broadcast to all connected passenger applications.')
add_normal('Key Components:')
add_bullet('GPS Module: Captures real-time location coordinates from the driver device.')
add_bullet('Socket.IO Client: Transmits location updates instantly to the backend.')
add_bullet('Tracking State Map: Maintains authoritative bus positions on the backend.')
add_bullet('WebView postMessage: Forwards updates to the embedded Leaflet.js map.')
add_normal('Challenges:')
add_bullet('GPS accuracy may vary in urban areas with signal obstruction.')
add_bullet('Continuous data transmission requires stable internet connectivity.')
add_bullet('Maintaining a single socket architecture to avoid duplicate connections.')

add_heading_custom('6.2 Driver Application Module', level=3)
add_normal('Objective: To enable drivers to manage trips, transmit GPS data, and trigger emergency alerts.')
add_normal('This module focuses on the driver-facing mobile application built with React Native and Expo. It provides secure login, route assignment, trip start and stop functionality, background GPS tracking, and an SOS emergency button. The driver application establishes a Socket.IO connection to the backend and emits location updates at regular intervals. It also displays route and bus details to assist the driver during travel.')
add_normal('Key Components:')
add_bullet('React Native with Expo: Cross-platform mobile framework.')
add_bullet('Background GPS Tracking: Captures location even when the app is backgrounded.')
add_bullet('Socket.IO Client: Sends real-time location and receives acknowledgments.')
add_bullet('SOS Button: Triggers emergency alerts broadcast to all connected clients.')
add_bullet('Trip Lifecycle Management: Handles start trip, stop trip, and route assignment.')
add_normal('Challenges:')
add_bullet('Requires stable internet connection for uninterrupted data transfer.')
add_bullet('Power consumption management for continuous GPS and socket operation.')
add_bullet('Ensuring accurate background location tracking on both Android and iOS.')

add_heading_custom('6.3 Backend API and Socket.IO Module', level=3)
add_normal('Objective: To handle data processing, real-time communication, and system logic through APIs and Socket.IO events.')
add_normal('This module is responsible for managing all backend operations of V-Bus. It receives real-time location data from the driver application through Socket.IO, processes it via the route-aware projection engine, and stores it in MongoDB. The backend maintains an authoritative tracking state map, manages stop progression logic, handles BUS_OFFLINE cleanup, and provides RESTful APIs for communication between the server, mobile applications, and web interface. It ensures secure access using JWT-based authentication and manages system functionalities such as bus tracking, route management, ETA calculation, and SOS event broadcasting.')
add_normal('Key Components:')
add_bullet('Node.js / Express.js Server: Processes requests, responses, and Socket.IO events.')
add_bullet('Socket.IO Server: Manages real-time bidirectional communication and room subscriptions.')
add_bullet('MongoDB with Mongoose: Stores system data and location history.')
add_bullet('Authentication System: Secures user access using JWT.')
add_bullet('Route-Aware Projection Engine: Projects bus location along route corridors.')
add_bullet('Tracking State Map: Maintains authoritative state of all active buses.')
add_bullet('Offline Cleanup System: Automatically removes stale bus data.')
add_normal('Challenges:')
add_bullet('Handling large amounts of real-time data efficiently.')
add_bullet('Ensuring security and preventing unauthorized access.')
add_bullet('Maintaining socket connection stability across diverse network conditions.')

add_heading_custom('6.4 Passenger Application Module', level=3)
add_normal('Objective: To provide a user-friendly interface for passengers to access real-time bus information.')
add_normal('This module is designed to allow passengers to interact with the system through a mobile application built with React Native. Users can track live bus locations, view route details, check estimated arrival times (ETA), and follow specific buses. The application communicates with the backend server through Socket.IO and forwards data to an embedded WebView rendering Leaflet.js maps over OpenStreetMap tiles. The BusContext architecture manages React state, socket subscriptions, and postMessage communication with the WebView.')
add_normal('Key Components:')
add_bullet('React Native with Expo: Provides cross-platform access.')
add_bullet('Embedded WebView: Renders interactive Leaflet.js maps.')
add_bullet('BusContext: Manages global state, socket lifecycle, and WebView communication.')
add_bullet('postMessage API: Enables React Native to WebView data exchange.')
add_bullet('Socket.IO Client: Receives real-time bus updates from the backend.')
add_bullet('Marker Replacement Strategy: Updates markers without map flicker.')
add_normal('Challenges:')
add_bullet('Requires continuous internet connectivity for live updates.')
add_bullet('Ensuring smooth WebView rendering across different devices.')
add_bullet('Managing socket reconnection and state synchronization efficiently.')

add_heading_custom('6.5 Route-Aware Projection Engine Module', level=3)
add_normal('Objective: To compute accurate ETA and project bus location along predefined route corridors.')
add_normal('This module is responsible for analyzing route geometries, projecting raw GPS coordinates onto the nearest point along the route corridor, and calculating the estimated time of arrival at upcoming stops. It ensures that the displayed bus location is visually consistent with the actual route path, rather than showing raw GPS points that may deviate due to road geometry or signal noise.')
add_normal('Key Components:')
add_bullet('Route Geometry Parser: Processes route corridor data from MongoDB.')
add_bullet('Projection Algorithm: Snaps raw GPS coordinates to the route path.')
add_bullet('Stop Progression Logic: Determines the next stop and arrival estimate.')
add_bullet('ETA Calculator: Computes estimated arrival based on speed and distance.')
add_normal('Challenges:')
add_bullet('Handling complex route geometries with multiple waypoints.')
add_bullet('Maintaining projection accuracy under GPS signal degradation.')
add_bullet('Balancing computational efficiency with real-time performance.')

add_heading_custom('6.6 Database and State Management Module', level=3)
add_normal('Objective: To store, manage, and retrieve system data efficiently.')
add_normal('This module is responsible for handling all data related to V-Bus. It stores information such as user details, bus data, route information, driver assignments, and real-time location history. The MongoDB database ensures that data is organized, easily accessible, and updated continuously based on system operations. The backend tracking state map complements the database by maintaining an in-memory authoritative view of all active buses for instantaneous query and broadcast.')
add_normal('Key Components:')
add_bullet('MongoDB: Stores structured and real-time data.')
add_bullet('Mongoose ODM: Defines schema for users, buses, routes, and locations.')
add_bullet('Query Processing: Retrieves and updates data efficiently.')
add_bullet('Tracking State Map: In-memory authoritative state of active buses.')
add_bullet('Offline Cleanup Logic: Removes stale records when buses disconnect.')
add_normal('Challenges:')
add_bullet('Managing large volumes of real-time data.')
add_bullet('Ensuring data consistency between in-memory state and database.')
add_bullet('Optimizing query performance for frequent location updates.')

doc.add_page_break()

# ===== CHAPTER 7: SYSTEM IMPLEMENTATION =====
add_heading_custom('CHAPTER 7', level=1)
add_heading_custom('SYSTEM IMPLEMENTATION', level=2)
add_normal('This chapter describes the practical implementation of V-Bus, including system development, integration of software components, and overall working process. The system combines driver mobile applications, backend services, passenger mobile interfaces, and embedded WebView maps to achieve real-time bus tracking and management.')

add_heading_custom('7.1 PROJECT STRUCTURE', level=3)
add_normal('The project is organized in a modular and scalable manner to separate concerns and facilitate ease of development and maintenance. Below is the directory structure of the project:')
add_normal('Figure – 7.1: Project Directory Structure')
add_normal('backend/')
add_bullet('src/routes/busRoutes.js: Express routes for bus operations, Socket.IO events, and tracking state management.')
add_bullet('src/models/User.js: Mongoose schema for user authentication and roles.')
add_bullet('src/services/ProjectionEngine.js: Route-aware projection and ETA calculation logic.')
add_bullet('server.js: Main entry point for Node.js, Express, and Socket.IO initialization.')
add_bullet('package.json: Backend dependencies including express, socket.io, mongoose, and dotenv.')
add_normal('passenger-app/')
add_bullet('App.js: Main application entry with navigation setup.')
add_bullet('BusContext.js: Global state management, socket lifecycle, and WebView communication.')
add_bullet('HomeScreen.js: Passenger home with nearby bus list, mini map, and real-time updates.')
add_bullet('FullMapScreen.js: Immersive full-screen map with follow-bus mode and route rendering.')
add_bullet('api/busApi.js: API client for RESTful backend communication.')
add_normal('driver-app/')
add_bullet('App.js: Driver application entry with authentication and navigation.')
add_bullet('DriverTrackingScreen.js: Background GPS tracking, trip lifecycle, and SOS button.')
add_bullet('RouteSelectionScreen.js: Route assignment and trip start/stop controls.')
add_normal('admin-web/')
add_bullet('src/App.jsx: React-based admin dashboard for bus, route, and driver management.')
add_bullet('public/: Static assets including favicon and icon sprites.')

add_heading_custom('7.2 DEVELOPMENT AND SUPPORTING TOOLS', level=3)
add_normal('The implementation of V-Bus utilizes various development tools and supporting libraries to ensure efficient system design, smooth integration, and reliable performance. These tools help in building the backend, frontend, mobile application, and map rendering components of the system.')

add_heading_custom('7.2.1 Development Tools', level=4)
add_normal('Table – 7.1: Development Tools')
add_table_grid([
    ('Visual Studio Code (VS Code)', 'Used as the primary code editor for developing backend and frontend applications.'),
    ('Node.js', 'Provides runtime environment for executing server-side JavaScript code.'),
    ('Express.js', 'Framework used to build RESTful APIs and serve Socket.IO events.'),
    ('React.js', 'Used to develop the web-based admin interface.'),
    ('React Native / Expo', 'Used to build cross-platform mobile applications for drivers and passengers.'),
    ('MongoDB', 'Database used to store system data such as buses, routes, users, and locations.'),
    ('Git', 'Version control system for managing source code.'),
], ['Tool', 'Purpose'])

add_heading_custom('7.2.2 Supporting Tools and Libraries', level=4)
add_normal('Table – 7.2: Supporting Tools and Libraries')
add_table_grid([
    ('OpenStreetMap', 'Provides free and open map tile data for Leaflet.js rendering.'),
    ('Leaflet.js', 'Open-source JavaScript library for interactive map visualization in WebView.'),
    ('Socket.IO', 'Enables real-time bidirectional communication between server and clients.'),
    ('Postman', 'Used for API testing and debugging.'),
    ('Mongoose', 'Object Data Modeling (ODM) for MongoDB.'),
    ('JWT (JSON Web Token)', 'Provides secure authentication and authorization.'),
    ('React Native WebView', 'Embedded browser component for rendering Leaflet.js maps.'),
    ('Expo Location', 'Provides background GPS tracking for driver application.'),
    ('dotenv', 'Manages environment variables for secure configuration.'),
    ('CORS', 'Enables cross-origin resource sharing for API security.'),
], ['Tool / Library', 'Functionality'])

doc.add_page_break()

# ===== CHAPTER 8: TESTING =====
add_heading_custom('CHAPTER 8', level=1)
add_heading_custom('TESTING', level=2)
add_normal('System testing is an important phase in the development of V-Bus, where the complete system is evaluated to ensure it meets the required specifications. Various testing methods are used to verify the functionality, performance, and reliability of the system.')

add_heading_custom('8.1 TESTING OBJECTIVES', level=3)
add_bullet('To verify that all modules of the system function correctly.')
add_bullet('To ensure seamless integration between driver application, backend, and passenger application components.')
add_bullet('To validate real-time tracking accuracy, socket synchronization, and system performance.')
add_bullet('To identify and fix errors or bugs before deployment.')
add_bullet('To ensure user-friendly and efficient system operation.')

add_heading_custom('8.2 TESTING STRATEGIES', level=3)

add_heading_custom('8.2.1 Unit Testing', level=4)
add_normal('Each module of the system is tested individually to ensure proper functionality. For example, GPS data collection, Socket.IO event emission, API responses, WebView postMessage communication, database operations, and route projection calculations are tested separately.')

add_heading_custom('8.2.2 Integration Testing', level=4)
add_normal('This testing ensures that all modules work together correctly. It verifies data flow from the driver application to the backend server via Socket.IO, and then to the passenger WebView interface. It also validates the BusContext state synchronization, marker replacement strategy, and BUS_OFFLINE cleanup events.')

add_heading_custom('8.2.3 Functional Testing', level=4)
add_normal('System features were validated against the functional requirements:')
add_normal('Table – 8.1: Functional Testing')
add_table_grid([
    ('Real-time bus tracking', 'Accurate live location displayed on WebView map', 'Pass'),
    ('Socket.IO synchronization', 'Instant updates across all connected clients', 'Pass'),
    ('ETA calculation', 'Correct arrival time shown at upcoming stops', 'Pass'),
    ('Route display', 'Proper route corridor and stops visible', 'Pass'),
    ('User login', 'Successful JWT-based authentication', 'Pass'),
    ('Driver trip lifecycle', 'Start/stop trip updates tracking state', 'Pass'),
    ('BUS_OFFLINE cleanup', 'Stale bus markers removed automatically', 'Pass'),
    ('SOS emergency alert', 'Emergency event broadcast to all clients', 'Pass'),
    ('Follow bus mode', 'Map viewport centers on selected bus', 'Pass'),
    ('Nearby stop detection', 'Nearest stops identified and displayed', 'Pass'),
], ['Feature', 'Expected Outcome', 'Result'])

add_heading_custom('8.2.4 Performance Testing', level=4)
add_normal('The system was tested for response time and real-time tracking performance.')
add_bullet('Average Tracking Delay: < 1–2 seconds')
add_bullet('Socket.IO Broadcast Latency: < 100 ms under stable network')
add_bullet('System Response Time: Fast response for user requests')
add_bullet('Data Processing: Efficient handling of real-time GPS and projection data')
add_bullet('App Performance: Smooth WebView map rendering without lag')

add_heading_custom('8.2.5 Usability Testing', level=4)
add_normal('Real users were asked to interact with the system through the mobile and web interfaces.')
add_bullet('User Feedback: Positive response regarding ease of use, map clarity, and interface design.')
add_bullet('Navigation: Simple and intuitive navigation across features.')
add_bullet('Accessibility: Easy access to tracking, routes, ETA, and nearby stops.')
add_bullet('User Experience: Smooth and user-friendly interaction with real-time updates.')

add_heading_custom('8.3 TEST CASES', level=3)
add_normal('Table – 8.2: Test Cases')
add_table_grid([
    ('TC001', 'Test real-time bus tracking', 'Driver app GPS active', 'Live bus marker displayed on passenger map', 'Pass'),
    ('TC002', 'Test ETA calculation', 'Bus approaching stop', 'Correct ETA shown in passenger app', 'Pass'),
    ('TC003', 'Test route corridor display', 'Select route', 'Route polyline and stops visible on map', 'Pass'),
    ('TC004', 'Test user login', 'Valid JWT credentials', 'Login successful and token stored', 'Pass'),
    ('TC005', 'Test socket data update', 'Driver location change', 'Real-time marker update reflected', 'Pass'),
    ('TC006', 'Test BUS_OFFLINE cleanup', 'Driver stops trip', 'Bus marker removed from all maps', 'Pass'),
    ('TC007', 'Test SOS emergency alert', 'Driver presses SOS', 'Alert broadcast to all connected clients', 'Pass'),
    ('TC008', 'Test follow bus mode', 'Enable follow mode', 'Map viewport centers on selected bus', 'Pass'),
    ('TC009', 'Test nearby stops', 'Passenger location', 'Nearest route stops listed', 'Pass'),
    ('TC010', 'Test admin route management', 'Create route', 'Route stored in MongoDB and available', 'Pass'),
], ['Test Case ID', 'Description', 'Input', 'Expected Output', 'Status'])

doc.add_page_break()

print("Completed through Chapter 8")
doc.save(r'w:\Final year project\V-Bus_Report.docx')
