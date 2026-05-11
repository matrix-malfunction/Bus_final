from docx import Document

doc = Document(r'w:\Final year project\Report.docx')

with open(r'w:\Final year project\report_extracted.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'
        f.write(f'[{style_name}] {para.text}\n')
    
    f.write('\n\n=== TABLES ===\n\n')
    for i, table in enumerate(doc.tables):
        f.write(f'--- Table {i+1} ---\n')
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            f.write('| ' + ' | '.join(cells) + ' |\n')
        f.write('\n')

print('Extraction complete')
