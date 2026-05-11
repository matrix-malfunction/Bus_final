from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document(r'w:\Final year project\V-Bus_Report.docx')

def set_para_format(para, alignment=None, bold=False, italic=False, font_size=12, font_name='Times New Roman', space_after=Pt(6)):
    para_format = para.paragraph_format
    if alignment:
        para.alignment = alignment
    para_format.space_after = space_after
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))
        return p
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(10))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(8))
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(6))
        return p

def add_normal(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Inches(0.5), bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_para_format(p, alignment=alignment, space_after=Pt(6))
    p.paragraph_format.first_line_indent = first_line_indent
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
    p.paragraph_format.first_line_indent = Inches(0)
    return p

def add_table_grid(rows_data, header_row):
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    for i, txt in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows_data:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            row.cells[i].text = txt
    return table

# ===== CHAPTER 2: LITERATURE REVIEW =====
add_heading_custom('CHAPTER 2', level=1)
add_heading_custom('LITERATURE REVIEW', level=2)

add_heading_custom('2.1 GPS-BASED BUS TRACKING SYSTEMS', level=3)
add_normal('Authors: P. Verma, S. Singh, and R. Kumar')
add_normal('Year: 2021')
add_normal('Interpretation: This study presents a GPS-based bus tracking system designed to provide real-time location updates of public transportation vehicles. The system utilizes GPS modules installed in buses to collect latitude and longitude data, which is then transmitted to a centralized server using mobile networks. Users can access this information through web or mobile applications to track bus positions and estimate arrival times. The research highlights the effectiveness of integrating GPS with web technologies to improve passenger convenience and reduce waiting time.')
add_normal('Pitfalls: The system heavily depends on stable internet connectivity and may experience delays in areas with poor network coverage. Additionally, it lacks advanced features such as predictive analytics, dynamic route optimization, and real-time bidirectional notifications, limiting its scalability and overall efficiency.')

add_heading_custom('2.2 REAL-TIME COMMUNICATION IN TRANSPORTATION', level=3)
add_normal('Authors: A. Sharma, K. Gupta, and M. Patel')
add_normal('Year: 2022')
add_normal('Interpretation: This research focuses on the use of event-driven real-time communication technologies for monitoring transportation systems. The proposed system integrates WebSocket-based protocols and GPS modules to continuously collect and transmit vehicle data such as location, speed, and status to a cloud-based platform. The study demonstrates how bidirectional communication enables seamless interaction between devices and centralized systems, allowing users and administrators to access real-time information through web and mobile applications. It also emphasizes improved operational efficiency, reduced delays, and better resource management in public transport systems.')
add_normal('Pitfalls: The system requires reliable network connectivity and low-latency infrastructure, which may not always be feasible in all environments. Additionally, issues related to data security, connection management, and scalability can affect long-term performance. The lack of an authoritative backend state and intelligent decision-making capabilities also limits the system’s effectiveness in complex urban scenarios.')

add_heading_custom('2.3 EVENT-DRIVEN VEHICLE MONITORING SYSTEMS', level=3)
add_normal('Authors: J. Lee, H. Kim, and S. Park')
add_normal('Year: 2020')
add_normal('Interpretation: This study presents an event-driven vehicle monitoring system that uses GPS and socket-based communication technologies to track the movement of vehicles continuously. The system collects data such as location, speed, and travel status, which is transmitted to a centralized server for processing and visualization. Users can access this information through web-based dashboards, enabling efficient monitoring and management of transportation systems. The research highlights the importance of real-time event propagation in improving operational control, reducing delays, and enhancing passenger satisfaction.')
add_normal('Pitfalls: The system faces challenges related to high-frequency data transmission, which can increase network load and operational costs. It also depends on accurate GPS signals, which may be affected in urban areas with signal obstructions. Additionally, the system lacks integration with advanced features such as route-aware projection, backend authoritative state management, and automated offline cleanup, limiting its overall effectiveness.')

add_heading_custom('2.4 SMART PUBLIC TRANSPORT SYSTEMS', level=3)
add_normal('Authors: R. Ahmed, L. Chen, and D. Williams')
add_normal('Year: 2023')
add_normal('Interpretation: This research explores smart public transport systems that integrate advanced technologies such as real-time communication, cloud computing, and mobile applications to improve the efficiency and reliability of urban transportation. The system enables real-time tracking, route management, digital passenger information services, and driver monitoring through a unified platform. It enhances communication between passengers and transport authorities, allowing users to access live updates, plan journeys effectively, and reduce waiting time. The study emphasizes the role of smart systems in building sustainable and intelligent urban mobility solutions.')
add_normal('Pitfalls: The implementation of such systems requires significant infrastructure investment and technical expertise. Data privacy and security concerns also arise due to continuous data collection and sharing. Additionally, integration with existing legacy transport systems can be complex, and the system’s performance may be affected by network reliability and scalability challenges. Many implementations also lack production stability and robust state management.')

add_heading_custom('2.5 WEBVIEW AND HYBRID MAP RENDERING', level=3)
add_normal('Authors: S. Gupta, N. Verma, and T. Reddy')
add_normal('Year: 2021')
add_normal('Interpretation: This study focuses on mobile-based bus tracking applications that leverage embedded WebView components to render interactive maps and provide real-time information to passengers through smartphones. The system uses GPS data integrated with hybrid mobile applications to display live bus locations, route details, and estimated arrival times. The embedded WebView communicates with the native layer through postMessage APIs, enabling seamless data exchange and map control. It enhances user convenience by enabling passengers to plan their journeys efficiently and reduces uncertainty in public transportation. The research highlights the importance of user-friendly interfaces, real-time updates, and cross-platform compatibility in improving the overall commuting experience.')
add_normal('Pitfalls: The effectiveness of the system depends on continuous internet connectivity and accurate GPS data. In areas with poor network coverage, the application may fail to provide timely updates. Additionally, many applications lack advanced features such as route-aware projection, backend authoritative tracking state, marker replacement strategies, and integration with centralized transport management systems, limiting their scalability and usability.')

add_heading_custom('2.6 SURVEY ON SMART TRANSPORT TECHNOLOGIES', level=3)
add_normal('Authors: K. Sharma, R. Iyer, and P. Nair')
add_normal('Year: 2022')
add_normal('Interpretation: This survey paper provides a comprehensive overview of various smart transport technologies used in modern transportation systems, including GPS tracking, real-time communication, cloud computing, hybrid mobile applications, and embedded map rendering. The study compares different approaches in terms of efficiency, scalability, cost-effectiveness, and real-time performance. It highlights how the integration of these technologies enables intelligent transportation systems that improve traffic management, reduce delays, and enhance passenger convenience. The research also emphasizes the growing importance of data-driven decision-making, backend authoritative architecture, and event-driven synchronization in optimizing transport operations and planning.')
add_normal('Pitfalls: The survey identifies common challenges such as high implementation costs, dependency on network infrastructure, and issues related to data security and privacy. It also points out the lack of standardization across different systems and limited interoperability between technologies. Furthermore, many existing solutions do not fully utilize advanced analytics, route-aware projection engines, or artificial intelligence, which restricts their ability to handle complex real-world transportation scenarios effectively.')

doc.add_page_break()

# ===== CHAPTER 3: EXISTING AND PROPOSED SYSTEM =====
add_heading_custom('CHAPTER 3', level=1)
add_heading_custom('EXISTING AND PROPOSED SYSTEM', level=2)

add_heading_custom('3.1 Existing System', level=3)
add_normal('Over the years, several systems have been developed to track and monitor public transportation. These systems mainly rely on basic GPS tracking, manual communication, or standalone mobile applications to provide limited information about bus locations. While they offer some level of convenience, they often fail to deliver accurate, real-time, and user-friendly solutions required for modern transportation systems.')

add_heading_custom('3.1.1 Manual Bus Tracking', level=4)
add_normal('Traditional public transport systems depend heavily on fixed schedules and manual inquiries. Passengers usually wait at bus stops without knowing the exact arrival time of buses. This approach is inefficient and unreliable, especially during traffic congestion or unexpected delays. It results in increased waiting time and inconvenience for daily commuters.')

add_heading_custom('3.1.2 GPS Tracking without Integration', level=4)
add_normal('Some systems use GPS devices installed in buses to track their location. However, these systems are often not integrated with user applications or centralized platforms. The collected data is not effectively utilized for passenger information or decision-making, limiting its usefulness. There is no real-time synchronization mechanism or event-driven update pipeline to propagate location changes to end users.')

add_heading_custom('3.1.3 Basic Mobile Tracking Applications', level=4)
add_normal('Certain mobile applications provide bus tracking features using GPS data. While they offer some level of real-time tracking, they are often limited in functionality. Many applications lack features such as accurate ETA prediction, route corridor visualization, follow-bus mode, real-time notifications, and automated offline cleanup, reducing their effectiveness for users.')

add_heading_custom('3.1.4 Limitations of Existing Systems', level=4)
add_normal('The existing systems suffer from several significant limitations that affect their efficiency and usability. Most of these systems lack real-time accuracy and reliability, making it difficult for passengers to depend on the provided information. There is limited integration between hardware devices, software applications, and end users, which results in poor coordination and ineffective data utilization. Additionally, many systems do not have a centralized management platform with authoritative backend state, making it challenging for transport authorities to monitor and control operations efficiently. The user experience is also affected due to incomplete features such as the absence of route-aware projection, WebView-based map rendering, Socket.IO synchronization, and automatic BUS_OFFLINE cleanup. Furthermore, the continued dependency on manual processes, outdated technologies such as SQLite and Flask, and the absence of event-driven architecture reduce the overall effectiveness and scalability of existing solutions.')

add_heading_custom('3.2 Proposed System', level=3)
add_normal('To overcome the limitations of existing systems, the proposed solution introduces V-Bus, a Real-Time Smart Bus Tracking and Passenger Information System that leverages Node.js, Express.js, Socket.IO, MongoDB, React Native, and OpenStreetMap to provide real-time tracking and efficient transport management. The system is designed to offer accurate, reliable, and user-friendly services for both passengers and transport authorities. It ensures seamless communication between buses, servers, and users through a centralized, event-driven platform.')

add_heading_custom('3.2.1 System Overview', level=4)
add_normal('The proposed system captures real-time location data from bus drivers through a React Native mobile application with background GPS tracking. This data is transmitted instantly to the backend server through Socket.IO events, where it is processed, stored, and synchronized. Users can access this information through the passenger mobile application, which embeds a WebView to render interactive Leaflet.js maps over OpenStreetMap tiles. Passengers can track buses, view routes, check estimated arrival times (ETA), and follow selected buses in real time. The system also allows administrators to monitor and manage buses, drivers, and routes efficiently through the web admin panel.')

add_heading_custom('3.2.2 Core Modules and Technologies', level=4)
add_bullet('Driver Application Module: Captures real-time location data from drivers using React Native with Expo and background GPS tracking.')
add_bullet('Backend Server Module: Handles data processing, Socket.IO synchronization, API communication, and database management using Node.js and Express.js.')
add_bullet('Passenger Application Module: Enables passengers to track buses, view routes, and receive real-time updates through React Native with an embedded WebView.')
add_bullet('Route-Aware Projection Engine: Computes ETA, projects bus location along route corridors, and manages stop progression logic.')
add_bullet('Database and State Management Module: Stores all relevant data in MongoDB and maintains an authoritative tracking state map on the backend.')
add_bullet('Admin Panel Module: Allows transport authorities to manage routes, buses, drivers, and monitor real-time operations.')

add_heading_custom('3.2.3 Key Features', level=4)
add_bullet('Real-Time Bus Tracking: Provides live location of buses using GPS and Socket.IO, enabling passengers to track movement accurately.')
add_bullet('ETA Prediction: Calculates estimated arrival time based on the route-aware projection engine and current bus location.')
add_bullet('User-Friendly Interface: Offers simple and intuitive mobile interfaces with embedded WebView map rendering for easy access.')
add_bullet('Centralized Management: Allows administrators to monitor and manage buses, routes, and drivers efficiently.')
add_bullet('Real-Time Notifications: Sends alerts about bus arrival, delays, and SOS emergencies to users.')
add_bullet('Follow Bus Mode: Keeps the map viewport centered on the selected bus for continuous observation.')
add_bullet('Offline Cleanup: Automatically removes stale bus data from the system when a bus goes offline.')
add_bullet('SOS Emergency Handling: Enables drivers to trigger emergency alerts that are broadcast to all connected clients.')
add_bullet('Scalable Architecture: Designed to handle large data volumes and support future system expansion.')

add_heading_custom('3.2.4 Advantages Over Existing Systems', level=4)
add_normal('Table – 3.1: Advantages Over Existing Systems')
add_table_grid([
    ('Real-Time Support', 'Limited or delayed tracking', 'Accurate real-time bus tracking with Socket.IO'),
    ('Integration', 'Partial or no integration', 'Fully integrated event-driven system with WebView'),
    ('User Experience', 'Basic and less interactive', 'User-friendly mobile interface with embedded maps'),
    ('ETA Accuracy', 'Not available or inaccurate', 'Accurate ETA via route-aware projection engine'),
    ('Management', 'Manual or semi-automated', 'Centralized backend authoritative state management'),
    ('Scalability', 'Limited expansion capability', 'Highly scalable Node.js and MongoDB architecture'),
    ('Notifications', 'Not supported', 'Real-time alerts, updates, and SOS handling'),
    ('Map Rendering', 'Static or third-party dependent', 'OpenStreetMap + Leaflet.js via WebView'),
    ('Offline Handling', 'Stale data remains visible', 'Automatic BUS_OFFLINE cleanup'),
], ['Criteria', 'Existing Systems', 'Proposed System'])

doc.add_page_break()

print("Completed through Chapter 3")
doc.save(r'w:\Final year project\V-Bus_Report.docx')
