from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document(r'w:\Final year project\V-Bus_Report.docx')

def set_para_format(para, alignment=None, bold=False, italic=False, font_size=12, font_name='Times New Roman', space_after=Pt(6)):
    para_format = para.paragraph_format
    if alignment:
        para.alignment = alignment
    para_format.space_after = space_after
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))
        return p
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(10))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(8))
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(6))
        return p

def add_normal(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Inches(0.5), bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_para_format(p, alignment=alignment, space_after=Pt(6))
    p.paragraph_format.first_line_indent = first_line_indent
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
    p.paragraph_format.first_line_indent = Inches(0)
    return p

# ===== CHAPTER 7: SYSTEM IMPLEMENTATION =====
add_heading_custom('CHAPTER 7', level=1)
add_heading_custom('SYSTEM IMPLEMENTATION', level=2)
add_normal('This chapter describes the practical implementation of V-Bus, including system development, integration of software components, and overall working process. The system combines driver mobile applications, backend services, passenger mobile interfaces, and embedded WebView maps to achieve real-time bus tracking and management.')

add_heading_custom('7.1 PROJECT STRUCTURE', level=3)
add_normal('The project is organized in a modular and scalable manner to separate concerns and facilitate ease of development and maintenance. Below is the directory structure of the project:')
add_normal('Figure – 7.1: Project Directory Structure')
add_normal('backend/')
add_bullet('src/routes/busRoutes.js: Express routes for bus operations, Socket.IO events, and tracking state management.')
add_bullet('src/models/User.js: Mongoose schema for user authentication and roles.')
add_bullet('src/services/ProjectionEngine.js: Route-aware projection and ETA calculation logic.')
add_bullet('server.js: Main entry point for Node.js, Express, and Socket.IO initialization.')
add_bullet('package.json: Backend dependencies including express, socket.io, mongoose, and dotenv.')
add_normal('passenger-app/')
add_bullet('App.js: Main application entry with navigation setup.')
add_bullet('BusContext.js: Global state management, socket lifecycle, and WebView communication.')
add_bullet('HomeScreen.js: Passenger home with nearby bus list, mini map, and real-time updates.')
add_bullet('FullMapScreen.js: Immersive full-screen map with follow-bus mode and route rendering.')
add_bullet('api/busApi.js: API client for RESTful backend communication.')
add_normal('driver-app/')
add_bullet('App.js: Driver application entry with authentication and navigation.')
add_bullet('DriverTrackingScreen.js: Background GPS tracking, trip lifecycle, and SOS button.')
add_bullet('RouteSelectionScreen.js: Route assignment and trip start/stop controls.')
add_normal('admin-web/')
add_bullet('src/App.jsx: React-based admin dashboard for bus, route, and driver management.')
add_bullet('public/: Static assets including favicon and icon sprites.')

add_heading_custom('7.2 DEVELOPMENT AND SUPPORTING TOOLS', level=3)
add_normal('The implementation of V-Bus utilizes various development tools and supporting libraries to ensure efficient system design, smooth integration, and reliable performance. These tools help in building the backend, frontend, mobile application, and map rendering components of the system.')

add_heading_custom('7.2.1 Development Tools', level=4)
add_normal('Table – 7.1: Development Tools')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Tool'
t.rows[0].cells[1].text = 'Purpose'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('Visual Studio Code (VS Code)', 'Used as the primary code editor for developing backend and frontend applications.'),
    ('Node.js', 'Provides runtime environment for executing server-side JavaScript code.'),
    ('Express.js', 'Framework used to build RESTful APIs and serve Socket.IO events.'),
    ('React.js', 'Used to develop the web-based admin interface.'),
    ('React Native / Expo', 'Used to build cross-platform mobile applications for drivers and passengers.'),
    ('MongoDB', 'Database used to store system data such as buses, routes, users, and locations.'),
    ('Git', 'Version control system for managing source code.'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]

add_heading_custom('7.2.2 Supporting Tools and Libraries', level=4)
add_normal('Table – 7.2: Supporting Tools and Libraries')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Tool / Library'
t.rows[0].cells[1].text = 'Functionality'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('OpenStreetMap', 'Provides free and open map tile data for Leaflet.js rendering.'),
    ('Leaflet.js', 'Open-source JavaScript library for interactive map visualization in WebView.'),
    ('Socket.IO', 'Enables real-time bidirectional communication between server and clients.'),
    ('Postman', 'Used for API testing and debugging.'),
    ('Mongoose', 'Object Data Modeling (ODM) for MongoDB.'),
    ('JWT (JSON Web Token)', 'Provides secure authentication and authorization.'),
    ('React Native WebView', 'Embedded browser component for rendering Leaflet.js maps.'),
    ('Expo Location', 'Provides background GPS tracking for driver application.'),
    ('dotenv', 'Manages environment variables for secure configuration.'),
    ('CORS', 'Enables cross-origin resource sharing for API security.'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]

doc.add_page_break()

# ===== CHAPTER 8: TESTING =====
add_heading_custom('CHAPTER 8', level=1)
add_heading_custom('TESTING', level=2)
add_normal('System testing is an important phase in the development of V-Bus, where the complete system is evaluated to ensure it meets the required specifications. Various testing methods are used to verify the functionality, performance, and reliability of the system.')

add_heading_custom('8.1 TESTING OBJECTIVES', level=3)
add_bullet('To verify that all modules of the system function correctly.')
add_bullet('To ensure seamless integration between driver application, backend, and passenger application components.')
add_bullet('To validate real-time tracking accuracy, socket synchronization, and system performance.')
add_bullet('To identify and fix errors or bugs before deployment.')
add_bullet('To ensure user-friendly and efficient system operation.')

add_heading_custom('8.2 TESTING STRATEGIES', level=3)

add_heading_custom('8.2.1 Unit Testing', level=4)
add_normal('Each module of the system is tested individually to ensure proper functionality. For example, GPS data collection, Socket.IO event emission, API responses, WebView postMessage communication, database operations, and route projection calculations are tested separately.')

add_heading_custom('8.2.2 Integration Testing', level=4)
add_normal('This testing ensures that all modules work together correctly. It verifies data flow from the driver application to the backend server via Socket.IO, and then to the passenger WebView interface. It also validates the BusContext state synchronization, marker replacement strategy, and BUS_OFFLINE cleanup events.')

add_heading_custom('8.2.3 Functional Testing', level=4)
add_normal('System features were validated against the functional requirements:')
add_normal('Table – 8.1: Functional Testing')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Feature'
t.rows[0].cells[1].text = 'Expected Outcome'
t.rows[0].cells[2].text = 'Result'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('Real-time bus tracking', 'Accurate live location displayed on WebView map', 'Pass'),
    ('Socket.IO synchronization', 'Instant updates across all connected clients', 'Pass'),
    ('ETA calculation', 'Correct arrival time shown at upcoming stops', 'Pass'),
    ('Route display', 'Proper route corridor and stops visible', 'Pass'),
    ('User login', 'Successful JWT-based authentication', 'Pass'),
    ('Driver trip lifecycle', 'Start/stop trip updates tracking state', 'Pass'),
    ('BUS_OFFLINE cleanup', 'Stale bus markers removed automatically', 'Pass'),
    ('SOS emergency alert', 'Emergency event broadcast to all clients', 'Pass'),
    ('Follow bus mode', 'Map viewport centers on selected bus', 'Pass'),
    ('Nearby stop detection', 'Nearest stops identified and displayed', 'Pass'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]
    row.cells[2].text = r[2]

add_heading_custom('8.2.4 Performance Testing', level=4)
add_normal('The system was tested for response time and real-time tracking performance.')
add_bullet('Average Tracking Delay: < 1–2 seconds')
add_bullet('Socket.IO Broadcast Latency: < 100 ms under stable network')
add_bullet('System Response Time: Fast response for user requests')
add_bullet('Data Processing: Efficient handling of real-time GPS and projection data')
add_bullet('App Performance: Smooth WebView map rendering without lag')

add_heading_custom('8.2.5 Usability Testing', level=4)
add_normal('Real users were asked to interact with the system through the mobile and web interfaces.')
add_bullet('User Feedback: Positive response regarding ease of use, map clarity, and interface design.')
add_bullet('Navigation: Simple and intuitive navigation across features.')
add_bullet('Accessibility: Easy access to tracking, routes, ETA, and nearby stops.')
add_bullet('User Experience: Smooth and user-friendly interaction with real-time updates.')

add_heading_custom('8.3 TEST CASES', level=3)
add_normal('Table – 8.2: Test Cases')
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Test Case ID'
t.rows[0].cells[1].text = 'Description'
t.rows[0].cells[2].text = 'Input'
t.rows[0].cells[3].text = 'Expected Output'
t.rows[0].cells[4].text = 'Status'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('TC001', 'Test real-time bus tracking', 'Driver app GPS active', 'Live bus marker displayed on passenger map', 'Pass'),
    ('TC002', 'Test ETA calculation', 'Bus approaching stop', 'Correct ETA shown in passenger app', 'Pass'),
    ('TC003', 'Test route corridor display', 'Select route', 'Route polyline and stops visible on map', 'Pass'),
    ('TC004', 'Test user login', 'Valid JWT credentials', 'Login successful and token stored', 'Pass'),
    ('TC005', 'Test socket data update', 'Driver location change', 'Real-time marker update reflected', 'Pass'),
    ('TC006', 'Test BUS_OFFLINE cleanup', 'Driver stops trip', 'Bus marker removed from all maps', 'Pass'),
    ('TC007', 'Test SOS emergency alert', 'Driver presses SOS', 'Alert broadcast to all connected clients', 'Pass'),
    ('TC008', 'Test follow bus mode', 'Enable follow mode', 'Map viewport centers on selected bus', 'Pass'),
    ('TC009', 'Test nearby stops', 'Passenger location', 'Nearest route stops listed', 'Pass'),
    ('TC010', 'Test admin route management', 'Create route', 'Route stored in MongoDB and available', 'Pass'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]
    row.cells[2].text = r[2]
    row.cells[3].text = r[3]
    row.cells[4].text = r[4]

doc.add_page_break()

# ===== CHAPTER 9: RESULTS AND PERFORMANCE =====
add_heading_custom('CHAPTER 9', level=1)
add_heading_custom('RESULTS AND PERFORMANCE', level=2)
add_normal('V-Bus was successfully implemented and tested to provide real-time tracking and efficient transport management. The system effectively collects live location data from the driver mobile application using background GPS tracking and updates it to the backend server through Socket.IO. This data is then synchronized across all connected passenger applications and displayed on embedded WebView maps, enabling accurate tracking of buses.')
add_normal('The system achieved high accuracy in location tracking under normal conditions, providing reliable real-time updates to users. The estimated time of arrival (ETA) feature produced accurate results based on the route-aware projection engine and current bus location. The average delay in updating the bus location was observed to be less than 1–2 seconds, ensuring smooth and responsive performance.')
add_normal('The application demonstrated stable performance across different devices, including smartphones and web browsers. Users reported a positive experience in terms of ease of use, clarity of map information, real-time synchronization, and overall system responsiveness. The backend system efficiently handled multiple concurrent socket connections and ensured consistent data processing without significant delays. The BUS_OFFLINE cleanup mechanism and SOS handling operated reliably, maintaining data consistency and safety.')
add_normal('Overall, the system met its objectives of providing accurate, fast, and user-friendly bus tracking services. It improved passenger convenience, reduced waiting time, enhanced operational visibility, and contributed to the efficiency of transport management systems.')

doc.add_page_break()

# ===== CHAPTER 10: CONCLUSION AND FUTURE ENHANCEMENTS =====
add_heading_custom('CHAPTER 10', level=1)
add_heading_custom('CONCLUSION AND FUTURE ENHANCEMENTS', level=2)

add_heading_custom('10.1 CONCLUSION', level=3)
add_normal('V-Bus provides an effective solution for real-time monitoring of public transportation. By integrating technologies such as GPS, Socket.IO, React Native, Node.js, Express.js, MongoDB, and OpenStreetMap, the system successfully enables passengers to track bus locations, view routes, estimate arrival times, and follow specific buses with accuracy. This enhances user convenience and reduces uncertainty in daily commuting.')
add_normal('The project demonstrates the importance of combining modern software architecture and real-time communication to build an efficient and scalable system. The use of React Native with Expo ensures cross-platform compatibility, while the Node.js and Express.js backend with Socket.IO provides instantaneous bidirectional data synchronization. The embedded WebView with Leaflet.js delivers immersive map visualization, and the route-aware projection engine ensures accurate ETA and corridor alignment. The system also supports centralized management, allowing transport authorities to monitor operations, respond to emergencies, and improve decision-making.')
add_normal('Overall, V-Bus meets its objectives by providing accurate, reliable, and user-friendly bus tracking services. It improves the efficiency of transportation systems and contributes to the development of smarter urban mobility solutions.')

add_heading_custom('10.2 FUTURE ENHANCEMENTS', level=3)
add_normal('To further improve the system, the following enhancements can be considered:')
add_bullet('Integration of AI for traffic prediction and dynamic route optimization.')
add_bullet('Support for multiple cities and large-scale deployment.')
add_bullet('Implementation of push notifications for real-time alerts.')
add_bullet('Integration with digital ticketing and payment systems.')
add_bullet('Advanced analytics for performance monitoring and reporting.')
add_bullet('Offline map caching for improved passenger experience in low-connectivity areas.')
add_bullet('Voice-based assistance and accessibility features for differently-abled passengers.')

doc.add_page_break()

print("Completed through Chapter 10")
doc.save(r'w:\Final year project\V-Bus_Report.docx')
