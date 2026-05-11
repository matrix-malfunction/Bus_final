from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document(r'w:\Final year project\V-Bus_Report.docx')

def set_para_format(para, alignment=None, bold=False, italic=False, font_size=12, font_name='Times New Roman', space_after=Pt(6)):
    para_format = para.paragraph_format
    if alignment:
        para.alignment = alignment
    para_format.space_after = space_after
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))
        return p
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(10))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(8))
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(6))
        return p

def add_normal(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Inches(0.5), bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_para_format(p, alignment=alignment, space_after=Pt(6))
    p.paragraph_format.first_line_indent = first_line_indent
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
    p.paragraph_format.first_line_indent = Inches(0)
    return p

# ===== CHAPTER 4: SYSTEM SPECIFICATION =====
add_heading_custom('CHAPTER 4', level=1)
add_heading_custom('SYSTEM SPECIFICATION', level=2)

add_heading_custom('4.1 HARDWARE REQUIREMENTS', level=3)
add_normal('The hardware resources are essential to ensure smooth operation, real-time processing, and optimal user experience. The primary hardware components include:')
add_normal('Table – 4.1: Hardware Requirements')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Component'
t.rows[0].cells[1].text = 'Minimum Requirement'
t.rows[0].cells[2].text = 'Recommended Requirement'
t.rows[0].cells[3].text = 'Purpose'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('Processor (CPU)', 'Intel Core i3 / Dual Core', 'Intel Core i5 / i7 or equivalent', 'Handles backend processing and system operations'),
    ('RAM', '4 GB', '8 GB or above', 'Supports application and server performance'),
    ('Storage', '50 GB HDD', '100 GB+ SSD', 'Stores database and application data'),
    ('GPS Device', 'Standard mobile GPS', 'High-accuracy GNSS chip', 'Captures real-time location coordinates from driver device'),
    ('Network Device', 'Basic WiFi Router', 'High-speed broadband router', 'Enables data transmission and connectivity'),
    ('Mobile Device', 'Android / iOS smartphone', 'Modern smartphone with GPS', 'Driver and passenger application host'),
    ('Power Supply', 'Standard charger', 'Reliable power source', 'Powers mobile devices during operation'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]
    row.cells[2].text = r[2]
    row.cells[3].text = r[3]

add_heading_custom('4.2 SOFTWARE REQUIREMENTS', level=3)
add_normal('The software components are essential for developing, managing, and running V-Bus efficiently. The following table outlines the minimum and recommended software requirements.')
add_normal('Table – 4.2: Software Requirements')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Category'
t.rows[0].cells[1].text = 'Minimum Requirement'
t.rows[0].cells[2].text = 'Recommended Tools/Frameworks'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('Operating System', 'Windows 10 / Linux / macOS', 'Windows 11 / Ubuntu 20.04+'),
    ('Programming Language', 'JavaScript (ES6+)', 'Node.js (Backend), JavaScript (Frontend)'),
    ('Backend Framework', 'Basic Node.js setup', 'Express.js with Socket.IO'),
    ('Frontend Development', 'HTML5, CSS3', 'React.js / React Native'),
    ('Mobile Development', 'Expo SDK', 'React Native with Expo'),
    ('Database', 'Local JSON storage', 'MongoDB with Mongoose ODM'),
    ('Map Rendering', 'Basic tile viewer', 'Leaflet.js with OpenStreetMap'),
    ('API Testing', 'Basic tools', 'Postman / Insomnia'),
    ('Development Tools', 'Any code editor', 'VS Code, Git'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]
    row.cells[2].text = r[2]

add_heading_custom('4.3 NETWORK REQUIREMENTS', level=3)
add_normal('A stable and reliable network connection is essential for real-time data transmission between the driver application, backend server, and passenger applications. The system depends on internet connectivity to send GPS data from the driver device to the backend and to provide live updates to users:')
add_normal('Table – 4.3: Network Requirements')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Requirement'
t.rows[0].cells[1].text = 'Minimum'
t.rows[0].cells[2].text = 'Recommended'
t.rows[0].cells[3].text = 'Purpose'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('Internet Speed', '2–5 Mbps', '10–20 Mbps or higher', 'Real-time data transmission and API communication'),
    ('Connectivity Type', 'WiFi / Mobile Data', 'High-speed Broadband / 4G/5G', 'Ensures uninterrupted communication'),
    ('Latency', '< 500 ms', '< 100 ms', 'Maintains low-latency Socket.IO event delivery'),
    ('Network Coverage', 'Limited coverage', 'Wide and stable coverage', 'Maintains consistent tracking accuracy'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]
    row.cells[2].text = r[2]
    row.cells[3].text = r[3]

add_heading_custom('4.4 SECURITY REQUIREMENTS', level=3)
add_normal('Security is an important aspect of V-Bus to protect user data and ensure safe communication between system components. The system implements various security measures to prevent unauthorized access and data breaches:')
add_normal('Table – 4.4: Security Requirements')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Security Feature'
t.rows[0].cells[1].text = 'Description'
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True
rows = [
    ('Encryption', 'SSL/TLS to protect user–system communication'),
    ('Authentication', 'Implements JWT-based authentication for secure login'),
    ('Authorization', 'Role-based access control for admin, drivers, and passengers'),
    ('Data Protection', 'Secures sensitive user and system data in MongoDB'),
    ('API Security', 'Protects endpoints from unauthorized access'),
    ('Socket Security', 'Validates socket connections and restricts event broadcasting'),
    ('Input Validation', 'Sanitizes all incoming requests to prevent injection attacks'),
]
for r in rows:
    row = t.add_row()
    row.cells[0].text = r[0]
    row.cells[1].text = r[1]

doc.add_page_break()

# ===== CHAPTER 5: SYSTEM DESIGN =====
add_heading_custom('CHAPTER 5', level=1)
add_heading_custom('SYSTEM DESIGN', level=2)

add_heading_custom('5.1 UML DIAGRAMS', level=3)
add_normal('Unified Modeling Language (UML) diagrams are used to represent the structure and behavior of a system in a visual format. These diagrams help in understanding the system architecture, data flow, and interaction between different components. In V-Bus, UML diagrams are used to illustrate how users, devices, and system modules interact to achieve real-time bus tracking and management.')

add_heading_custom('5.1.1 USE CASE DIAGRAM', level=4)
add_normal('Figure – 5.1: Use Case Diagram')
add_normal('The use case diagram represents the interaction between different actors and the system. In V-Bus, the primary actors include passengers, administrators, and drivers. Passengers use the system to track buses, view routes, check estimated arrival times, and follow specific buses. Administrators manage buses, routes, and drivers through the admin panel, while drivers update trip status, transmit GPS data, and trigger SOS alerts through the driver application.')
add_normal('The system processes user requests, retrieves real-time location data from the driver application via Socket.IO, and displays the information through the passenger mobile interface and web admin dashboard. This diagram helps in understanding the overall functionality of the system and the roles of each user in interacting with the system.')

add_heading_custom('5.1.2 ACTIVITY DIAGRAM', level=4)
add_normal('Figure – 5.2: Activity Diagram')
add_normal('The activity diagram illustrates the flow of operations involved in V-Bus. It represents the sequence of activities starting from user input to the final output displayed on the system. The process begins when the passenger opens the application and requests bus tracking information.')
add_normal('The system then retrieves real-time location data from the driver mobile application through background GPS tracking. This data is transmitted to the backend server via Socket.IO events, where it is processed, projected along route corridors, and stored in the MongoDB database. The server calculates the estimated time of arrival (ETA) and fetches relevant route information.')
add_normal('Once the data is processed, the system sends the updated information to the passenger application through Socket.IO broadcasts, where the live bus location and details are displayed on the WebView map. If the user requests additional information such as route details, nearby stops, or SOS alerts, the system processes those requests accordingly. The activity diagram helps in understanding the step-by-step workflow and decision-making process within the system.')

add_heading_custom('5.1.3 SEQUENCE DIAGRAM', level=4)
add_normal('Figure – 5.3: Sequence Diagram')
add_normal('The sequence diagram illustrates the interaction between different components of V-Bus in a time-ordered manner. It shows how data flows between the passenger, driver application, backend server, MongoDB database, and WebView map during the bus tracking process.')
add_normal('The driver application captures GPS coordinates and emits a driver_location_update event through Socket.IO. The backend receives this event, updates the authoritative tracking state map, projects the location along the assigned route corridor, and broadcasts a bus_location_update event to all connected passenger clients. The passenger application receives this event, forwards the data to the WebView via postMessage, and the Leaflet.js map updates the bus marker position. When the driver stops the trip or goes offline, the backend emits a BUS_OFFLINE event, triggering marker removal and cleanup in the WebView.')

add_heading_custom('5.1.4 CLASS DIAGRAM', level=4)
add_normal('Figure – 5.4: Class Diagram')
add_normal('The class diagram represents the static structure of V-Bus by illustrating the classes, their attributes, methods, and relationships. It helps in understanding how different components of the system are organized and how they interact with each other.')
add_normal('Key classes include the Bus class (with attributes such as busId, routeId, currentLocation, and status), the Driver class (with driverId, assignedBus, and tripStatus), the Route class (with routeId, stops, and corridorGeometry), the TrackingState class (maintaining the authoritative state map of all active buses), the SocketManager class (handling connection, event registration, and room management), and the BusContext class (managing React state, socket subscriptions, and WebView communication in the passenger application).')

add_heading_custom('5.1.5 ARCHITECTURE DIAGRAM', level=4)
add_normal('Figure – 5.5: Architecture Diagram')
add_normal('The architecture diagram illustrates the overall structure of V-Bus and the interaction between its major components. It provides a high-level view of how data flows from the driver application layer to the user interface.')
add_normal('The system begins with the driver mobile application built in React Native with Expo, which continuously collects real-time location data using the device GPS. This data is transmitted over the internet to the Node.js and Express.js backend server through Socket.IO events. The backend processes the incoming data, updates the tracking state map, stores it in the MongoDB database, and performs necessary computations such as ETA calculation and route corridor projection.')
add_normal('The processed data is then broadcast to all connected passenger applications through Socket.IO. The passenger application, built with React Native, hosts an embedded WebView that renders a Leaflet.js map over OpenStreetMap tiles. Communication between the React Native layer and the WebView occurs via postMessage, allowing real-time marker updates, route rendering, popup information, and offline cleanup. Passengers can access real-time bus location, route details, nearby stops, and notifications, while administrators can manage buses, routes, and drivers through the web admin panel. The architecture ensures smooth communication between all components, enabling efficient and real-time system performance.')

doc.add_page_break()

# ===== CHAPTER 6: SYSTEM MODULES =====
add_heading_custom('CHAPTER 6', level=1)
add_heading_custom('SYSTEM MODULES', level=2)
add_normal('V-Bus is composed of multiple interconnected modules that work together to provide real-time tracking and efficient transport management. Each module performs a specific function and contributes to the overall system performance. These modules ensure smooth data collection, processing, and user interaction across the platform.')
add_normal('Figure – 6.1: System Modules')

add_heading_custom('6.1 Real-Time Tracking Module', level=3)
add_normal('Objective: To provide accurate real-time tracking of buses using GPS and Socket.IO.')
add_normal('This module is responsible for tracking the live location of buses using GPS technology. It collects location data such as latitude and longitude from the driver mobile application and updates it continuously through Socket.IO events. The data is transmitted to the backend server, where it is processed by the route-aware projection engine and broadcast to all connected passenger applications.')
add_normal('Key Components:')
add_bullet('GPS Module: Captures real-time location coordinates from the driver device.')
add_bullet('Socket.IO Client: Transmits location updates instantly to the backend.')
add_bullet('Tracking State Map: Maintains authoritative bus positions on the backend.')
add_bullet('WebView postMessage: Forwards updates to the embedded Leaflet.js map.')
add_normal('Challenges:')
add_bullet('GPS accuracy may vary in urban areas with signal obstruction.')
add_bullet('Continuous data transmission requires stable internet connectivity.')
add_bullet('Maintaining a single socket architecture to avoid duplicate connections.')

add_heading_custom('6.2 Driver Application Module', level=3)
add_normal('Objective: To enable drivers to manage trips, transmit GPS data, and trigger emergency alerts.')
add_normal('This module focuses on the driver-facing mobile application built with React Native and Expo. It provides secure login, route assignment, trip start and stop functionality, background GPS tracking, and an SOS emergency button. The driver application establishes a Socket.IO connection to the backend and emits location updates at regular intervals. It also displays route and bus details to assist the driver during travel.')
add_normal('Key Components:')
add_bullet('React Native with Expo: Cross-platform mobile framework.')
add_bullet('Background GPS Tracking: Captures location even when the app is backgrounded.')
add_bullet('Socket.IO Client: Sends real-time location and receives acknowledgments.')
add_bullet('SOS Button: Triggers emergency alerts broadcast to all connected clients.')
add_bullet('Trip Lifecycle Management: Handles start trip, stop trip, and route assignment.')
add_normal('Challenges:')
add_bullet('Requires stable internet connection for uninterrupted data transfer.')
add_bullet('Power consumption management for continuous GPS and socket operation.')
add_bullet('Ensuring accurate background location tracking on both Android and iOS.')

add_heading_custom('6.3 Backend API and Socket.IO Module', level=3)
add_normal('Objective: To handle data processing, real-time communication, and system logic through APIs and Socket.IO events.')
add_normal('This module is responsible for managing all backend operations of V-Bus. It receives real-time location data from the driver application through Socket.IO, processes it via the route-aware projection engine, and stores it in MongoDB. The backend maintains an authoritative tracking state map, manages stop progression logic, handles BUS_OFFLINE cleanup, and provides RESTful APIs for communication between the server, mobile applications, and web interface. It ensures secure access using JWT-based authentication and manages system functionalities such as bus tracking, route management, ETA calculation, and SOS event broadcasting.')
add_normal('Key Components:')
add_bullet('Node.js / Express.js Server: Processes requests, responses, and Socket.IO events.')
add_bullet('Socket.IO Server: Manages real-time bidirectional communication and room subscriptions.')
add_bullet('MongoDB with Mongoose: Stores system data and location history.')
add_bullet('Authentication System: Secures user access using JWT.')
add_bullet('Route-Aware Projection Engine: Projects bus location along route corridors.')
add_bullet('Tracking State Map: Maintains authoritative state of all active buses.')
add_bullet('Offline Cleanup System: Automatically removes stale bus data.')
add_normal('Challenges:')
add_bullet('Handling large amounts of real-time data efficiently.')
add_bullet('Ensuring security and preventing unauthorized access.')
add_bullet('Maintaining socket connection stability across diverse network conditions.')

add_heading_custom('6.4 Passenger Application Module', level=3)
add_normal('Objective: To provide a user-friendly interface for passengers to access real-time bus information.')
add_normal('This module is designed to allow passengers to interact with the system through a mobile application built with React Native. Users can track live bus locations, view route details, check estimated arrival times (ETA), and follow specific buses. The application communicates with the backend server through Socket.IO and forwards data to an embedded WebView rendering Leaflet.js maps over OpenStreetMap tiles. The BusContext architecture manages React state, socket subscriptions, and postMessage communication with the WebView.')
add_normal('Key Components:')
add_bullet('React Native with Expo: Provides cross-platform access.')
add_bullet('Embedded WebView: Renders interactive Leaflet.js maps.')
add_bullet('BusContext: Manages global state, socket lifecycle, and WebView communication.')
add_bullet('postMessage API: Enables React Native to WebView data exchange.')
add_bullet('Socket.IO Client: Receives real-time bus updates from the backend.')
add_bullet('Marker Replacement Strategy: Updates markers without map flicker.')
add_normal('Challenges:')
add_bullet('Requires continuous internet connectivity for live updates.')
add_bullet('Ensuring smooth WebView rendering across different devices.')
add_bullet('Managing socket reconnection and state synchronization efficiently.')

add_heading_custom('6.5 Route-Aware Projection Engine Module', level=3)
add_normal('Objective: To compute accurate ETA and project bus location along predefined route corridors.')
add_normal('This module is responsible for analyzing route geometries, projecting raw GPS coordinates onto the nearest point along the route corridor, and calculating the estimated time of arrival at upcoming stops. It ensures that the displayed bus location is visually consistent with the actual route path, rather than showing raw GPS points that may deviate due to road geometry or signal noise.')
add_normal('Key Components:')
add_bullet('Route Geometry Parser: Processes route corridor data from MongoDB.')
add_bullet('Projection Algorithm: Snaps raw GPS coordinates to the route path.')
add_bullet('Stop Progression Logic: Determines the next stop and arrival estimate.')
add_bullet('ETA Calculator: Computes estimated arrival based on speed and distance.')
add_normal('Challenges:')
add_bullet('Handling complex route geometries with multiple waypoints.')
add_bullet('Maintaining projection accuracy under GPS signal degradation.')
add_bullet('Balancing computational efficiency with real-time performance.')

add_heading_custom('6.6 Database and State Management Module', level=3)
add_normal('Objective: To store, manage, and retrieve system data efficiently.')
add_normal('This module is responsible for handling all data related to V-Bus. It stores information such as user details, bus data, route information, driver assignments, and real-time location history. The MongoDB database ensures that data is organized, easily accessible, and updated continuously based on system operations. The backend tracking state map complements the database by maintaining an in-memory authoritative view of all active buses for instantaneous query and broadcast.')
add_normal('Key Components:')
add_bullet('MongoDB: Stores structured and real-time data.')
add_bullet('Mongoose ODM: Defines schema for users, buses, routes, and locations.')
add_bullet('Query Processing: Retrieves and updates data efficiently.')
add_bullet('Tracking State Map: In-memory authoritative state of active buses.')
add_bullet('Offline Cleanup Logic: Removes stale records when buses disconnect.')
add_normal('Challenges:')
add_bullet('Managing large volumes of real-time data.')
add_bullet('Ensuring data consistency between in-memory state and database.')
add_bullet('Optimizing query performance for frequent location updates.')

doc.add_page_break()

print("Completed through Chapter 6")
doc.save(r'w:\Final year project\V-Bus_Report.docx')
