from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document(r'w:\Final year project\V-Bus_Report.docx')

def set_para_format(para, alignment=None, bold=False, italic=False, font_size=12, font_name='Times New Roman', space_after=Pt(6)):
    para_format = para.paragraph_format
    if alignment:
        para.alignment = alignment
    para_format.space_after = space_after
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))
        return p
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(10))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(8))
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(6))
        return p

def add_normal(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Inches(0.5), bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_para_format(p, alignment=alignment, space_after=Pt(6))
    p.paragraph_format.first_line_indent = first_line_indent
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
    p.paragraph_format.first_line_indent = Inches(0)
    return p

def add_table_grid(rows_data, header_row):
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    for i, txt in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = txt
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows_data:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            row.cells[i].text = txt
    return table

# ===== CHAPTER 4: SYSTEM SPECIFICATION =====
add_heading_custom('CHAPTER 4', level=1)
add_heading_custom('SYSTEM SPECIFICATION', level=2)

add_heading_custom('4.1 HARDWARE REQUIREMENTS', level=3)
add_normal('The hardware resources are essential to ensure smooth operation, real-time processing, and optimal user experience. The primary hardware components include:')
add_normal('Table – 4.1: Hardware Requirements')
add_table_grid([
    ('Processor (CPU)', 'Intel Core i3 / Dual Core', 'Intel Core i5 / i7 or equivalent', 'Handles backend processing and system operations'),
    ('RAM', '4 GB', '8 GB or above', 'Supports application and server performance'),
    ('Storage', '50 GB HDD', '100 GB+ SSD', 'Stores database and application data'),
    ('GPS Device', 'Standard mobile GPS', 'High-accuracy GNSS chip', 'Captures real-time location coordinates from driver device'),
    ('Network Device', 'Basic WiFi Router', 'High-speed broadband router', 'Enables data transmission and connectivity'),
    ('Mobile Device', 'Android / iOS smartphone', 'Modern smartphone with GPS', 'Driver and passenger application host'),
    ('Power Supply', 'Standard charger', 'Reliable power source', 'Powers mobile devices during operation'),
], ['Component', 'Minimum Requirement', 'Recommended Requirement', 'Purpose'])

add_heading_custom('4.2 SOFTWARE REQUIREMENTS', level=3)
add_normal('The software components are essential for developing, managing, and running V-Bus efficiently. The following table outlines the minimum and recommended software requirements.')
add_normal('Table – 4.2: Software Requirements')
add_table_grid([
    ('Operating System', 'Windows 10 / Linux / macOS', 'Windows 11 / Ubuntu 20.04+'),
    ('Programming Language', 'JavaScript (ES6+)', 'Node.js (Backend), JavaScript (Frontend)'),
    ('Backend Framework', 'Basic Node.js setup', 'Express.js with Socket.IO'),
    ('Frontend Development', 'HTML5, CSS3', 'React.js / React Native'),
    ('Mobile Development', 'Expo SDK', 'React Native with Expo'),
    ('Database', 'Local JSON storage', 'MongoDB with Mongoose ODM'),
    ('Map Rendering', 'Basic tile viewer', 'Leaflet.js with OpenStreetMap'),
    ('API Testing', 'Basic tools', 'Postman / Insomnia'),
    ('Development Tools', 'Any code editor', 'VS Code, Git'),
], ['Category', 'Minimum Requirement', 'Recommended Tools/Frameworks'])

add_heading_custom('4.3 NETWORK REQUIREMENTS', level=3)
add_normal('A stable and reliable network connection is essential for real-time data transmission between the driver application, backend server, and passenger applications. The system depends on internet connectivity to send GPS data from the driver device to the backend and to provide live updates to users:')
add_normal('Table – 4.3: Network Requirements')
add_table_grid([
    ('Internet Speed', '2–5 Mbps', '10–20 Mbps or higher', 'Real-time data transmission and API communication'),
    ('Connectivity Type', 'WiFi / Mobile Data', 'High-speed Broadband / 4G/5G', 'Ensures uninterrupted communication'),
    ('Latency', '< 500 ms', '< 100 ms', 'Maintains low-latency Socket.IO event delivery'),
    ('Network Coverage', 'Limited coverage', 'Wide and stable coverage', 'Maintains consistent tracking accuracy'),
], ['Requirement', 'Minimum', 'Recommended', 'Purpose'])

add_heading_custom('4.4 SECURITY REQUIREMENTS', level=3)
add_normal('Security is an important aspect of V-Bus to protect user data and ensure safe communication between system components. The system implements various security measures to prevent unauthorized access and data breaches:')
add_normal('Table – 4.4: Security Requirements')
add_table_grid([
    ('Encryption', 'SSL/TLS to protect user–system communication'),
    ('Authentication', 'Implements JWT-based authentication for secure login'),
    ('Authorization', 'Role-based access control for admin, drivers, and passengers'),
    ('Data Protection', 'Secures sensitive user and system data in MongoDB'),
    ('API Security', 'Protects endpoints from unauthorized access'),
    ('Socket Security', 'Validates socket connections and restricts event broadcasting'),
    ('Input Validation', 'Sanitizes all incoming requests to prevent injection attacks'),
], ['Security Feature', 'Description'])

doc.add_page_break()

# ===== CHAPTER 5: SYSTEM DESIGN =====
add_heading_custom('CHAPTER 5', level=1)
add_heading_custom('SYSTEM DESIGN', level=2)

add_heading_custom('5.1 UML DIAGRAMS', level=3)
add_normal('Unified Modeling Language (UML) diagrams are used to represent the structure and behavior of a system in a visual format. These diagrams help in understanding the system architecture, data flow, and interaction between different components. In V-Bus, UML diagrams are used to illustrate how users, devices, and system modules interact to achieve real-time bus tracking and management.')

add_heading_custom('5.1.1 USE CASE DIAGRAM', level=4)
add_normal('Figure – 5.1: Use Case Diagram')
add_normal('The use case diagram represents the interaction between different actors and the system. In V-Bus, the primary actors include passengers, administrators, and drivers. Passengers use the system to track buses, view routes, check estimated arrival times, and follow specific buses. Administrators manage buses, routes, and drivers through the admin panel, while drivers update trip status, transmit GPS data, and trigger SOS alerts through the driver application.')
add_normal('The system processes user requests, retrieves real-time location data from the driver application via Socket.IO, and displays the information through the passenger mobile interface and web admin dashboard. This diagram helps in understanding the overall functionality of the system and the roles of each user in interacting with the system.')

add_heading_custom('5.1.2 ACTIVITY DIAGRAM', level=4)
add_normal('Figure – 5.2: Activity Diagram')
add_normal('The activity diagram illustrates the flow of operations involved in V-Bus. It represents the sequence of activities starting from user input to the final output displayed on the system. The process begins when the passenger opens the application and requests bus tracking information.')
add_normal('The system then retrieves real-time location data from the driver mobile application through background GPS tracking. This data is transmitted to the backend server via Socket.IO events, where it is processed, projected along route corridors, and stored in the MongoDB database. The server calculates the estimated time of arrival (ETA) and fetches relevant route information.')
add_normal('Once the data is processed, the system sends the updated information to the passenger application through Socket.IO broadcasts, where the live bus location and details are displayed on the WebView map. If the user requests additional information such as route details, nearby stops, or SOS alerts, the system processes those requests accordingly. The activity diagram helps in understanding the step-by-step workflow and decision-making process within the system.')

add_heading_custom('5.1.3 SEQUENCE DIAGRAM', level=4)
add_normal('Figure – 5.3: Sequence Diagram')
add_normal('The sequence diagram illustrates the interaction between different components of V-Bus in a time-ordered manner. It shows how data flows between the passenger, driver application, backend server, MongoDB database, and WebView map during the bus tracking process.')
add_normal('The driver application captures GPS coordinates and emits a driver_location_update event through Socket.IO. The backend receives this event, updates the authoritative tracking state map, projects the location along the assigned route corridor, and broadcasts a bus_location_update event to all connected passenger clients. The passenger application receives this event, forwards the data to the WebView via postMessage, and the Leaflet.js map updates the bus marker position. When the driver stops the trip or goes offline, the backend emits a BUS_OFFLINE event, triggering marker removal and cleanup in the WebView.')

add_heading_custom('5.1.4 CLASS DIAGRAM', level=4)
add_normal('Figure – 5.4: Class Diagram')
add_normal('The class diagram represents the static structure of V-Bus by illustrating the classes, their attributes, methods, and relationships. It helps in understanding how different components of the system are organized and how they interact with each other.')
add_normal('Key classes include the Bus class (with attributes such as busId, routeId, currentLocation, and status), the Driver class (with driverId, assignedBus, and tripStatus), the Route class (with routeId, stops, and corridorGeometry), the TrackingState class (maintaining the authoritative state map of all active buses), the SocketManager class (handling connection, event registration, and room management), and the BusContext class (managing React state, socket subscriptions, and WebView communication in the passenger application).')

add_heading_custom('5.1.5 ARCHITECTURE DIAGRAM', level=4)
add_normal('Figure – 5.5: Architecture Diagram')
add_normal('The architecture diagram illustrates the overall structure of V-Bus and the interaction between its major components. It provides a high-level view of how data flows from the driver application layer to the user interface.')
add_normal('The system begins with the driver mobile application built in React Native with Expo, which continuously collects real-time location data using the device GPS. This data is transmitted over the internet to the Node.js and Express.js backend server through Socket.IO events. The backend processes the incoming data, updates the tracking state map, stores it in the MongoDB database, and performs necessary computations such as ETA calculation and route corridor projection.')
add_normal('The processed data is then broadcast to all connected passenger applications through Socket.IO. The passenger application, built with React Native, hosts an embedded WebView that renders a Leaflet.js map over OpenStreetMap tiles. Communication between the React Native layer and the WebView occurs via postMessage, allowing real-time marker updates, route rendering, popup information, and offline cleanup. Passengers can access real-time bus location, route details, nearby stops, and notifications, while administrators can manage buses, routes, and drivers through the web admin panel. The architecture ensures smooth communication between all components, enabling efficient and real-time system performance.')

doc.add_page_break()

print("Completed through Chapter 5")
doc.save(r'w:\Final year project\V-Bus_Report.docx')
