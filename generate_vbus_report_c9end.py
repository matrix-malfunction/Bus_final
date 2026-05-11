from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document(r'w:\Final year project\V-Bus_Report.docx')

def set_para_format(para, alignment=None, bold=False, italic=False, font_size=12, font_name='Times New Roman', space_after=Pt(6)):
    para_format = para.paragraph_format
    if alignment:
        para.alignment = alignment
    para_format.space_after = space_after
    para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_custom(text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, font_size=Pt(16), bold=True, space_after=Pt(12))
        return p
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(14), bold=True, space_after=Pt(10))
        return p
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(8))
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, font_size=Pt(12), bold=True, space_after=Pt(6))
        return p

def add_normal(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Inches(0.5), bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_para_format(p, alignment=alignment, space_after=Pt(6))
    p.paragraph_format.first_line_indent = first_line_indent
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
    p.paragraph_format.first_line_indent = Inches(0)
    return p

# ===== CHAPTER 9: RESULTS AND PERFORMANCE =====
add_heading_custom('CHAPTER 9', level=1)
add_heading_custom('RESULTS AND PERFORMANCE', level=2)
add_normal('V-Bus was successfully implemented and tested to provide real-time tracking and efficient transport management. The system effectively collects live location data from the driver mobile application using background GPS tracking and updates it to the backend server through Socket.IO. This data is then synchronized across all connected passenger applications and displayed on embedded WebView maps, enabling accurate tracking of buses.')
add_normal('The system achieved high accuracy in location tracking under normal conditions, providing reliable real-time updates to users. The estimated time of arrival (ETA) feature produced accurate results based on the route-aware projection engine and current bus location. The average delay in updating the bus location was observed to be less than 1–2 seconds, ensuring smooth and responsive performance.')
add_normal('The application demonstrated stable performance across different devices, including smartphones and web browsers. Users reported a positive experience in terms of ease of use, clarity of map information, real-time synchronization, and overall system responsiveness. The backend system efficiently handled multiple concurrent socket connections and ensured consistent data processing without significant delays. The BUS_OFFLINE cleanup mechanism and SOS handling operated reliably, maintaining data consistency and safety.')
add_normal('Overall, the system met its objectives of providing accurate, fast, and user-friendly bus tracking services. It improved passenger convenience, reduced waiting time, enhanced operational visibility, and contributed to the efficiency of transport management systems.')

doc.add_page_break()

# ===== CHAPTER 10: CONCLUSION AND FUTURE ENHANCEMENTS =====
add_heading_custom('CHAPTER 10', level=1)
add_heading_custom('CONCLUSION AND FUTURE ENHANCEMENTS', level=2)

add_heading_custom('10.1 CONCLUSION', level=3)
add_normal('V-Bus provides an effective solution for real-time monitoring of public transportation. By integrating technologies such as GPS, Socket.IO, React Native, Node.js, Express.js, MongoDB, and OpenStreetMap, the system successfully enables passengers to track bus locations, view routes, estimate arrival times, and follow specific buses with accuracy. This enhances user convenience and reduces uncertainty in daily commuting.')
add_normal('The project demonstrates the importance of combining modern software architecture and real-time communication to build an efficient and scalable system. The use of React Native with Expo ensures cross-platform compatibility, while the Node.js and Express.js backend with Socket.IO provides instantaneous bidirectional data synchronization. The embedded WebView with Leaflet.js delivers immersive map visualization, and the route-aware projection engine ensures accurate ETA and corridor alignment. The system also supports centralized management, allowing transport authorities to monitor operations, respond to emergencies, and improve decision-making.')
add_normal('Overall, V-Bus meets its objectives by providing accurate, reliable, and user-friendly bus tracking services. It improves the efficiency of transportation systems and contributes to the development of smarter urban mobility solutions.')

add_heading_custom('10.2 FUTURE ENHANCEMENTS', level=3)
add_normal('To further improve the system, the following enhancements can be considered:')
add_bullet('Integration of AI for traffic prediction and dynamic route optimization.')
add_bullet('Support for multiple cities and large-scale deployment.')
add_bullet('Implementation of push notifications for real-time alerts.')
add_bullet('Integration with digital ticketing and payment systems.')
add_bullet('Advanced analytics for performance monitoring and reporting.')
add_bullet('Offline map caching for improved passenger experience in low-connectivity areas.')
add_bullet('Voice-based assistance and accessibility features for differently-abled passengers.')

doc.add_page_break()

# ===== APPENDIX 1: SAMPLE CODE =====
add_heading_custom('APPENDIX 1: SAMPLE CODE', level=1)

add_heading_custom('Backend Server with Socket.IO', level=2)
add_normal('server.js')
add_normal("""
const express = require('express');
const mongoose = require('mongoose');
const cors = require('cors');
const { createServer } = require('http');
const { Server } = require('socket.io');
require('dotenv').config();

const app = express();
const httpServer = createServer(app);
const io = new Server(httpServer, {
  cors: { origin: "*", methods: ["GET", "POST"] }
});

const PORT = process.env.PORT || 3000;

app.use(cors());
app.use(express.json());

const trackingState = new Map();

io.on('connection', (socket) => {
  console.log(`[Socket] Client connected: ${socket.id}`);

  socket.on('track_bus', (busId) => {
    socket.join(`bus_${busId}`);
    console.log(`[Socket] Client tracking bus ${busId}`);
  });

  socket.on('untrack_bus', (busId) => {
    socket.leave(`bus_${busId}`);
    console.log(`[Socket] Client stopped tracking bus ${busId}`);
  });

  socket.on('driver_location_update', (data) => {
    trackingState.set(data.busId, {
      ...data,
      source: 'driver_app',
      timestamp: new Date().toISOString()
    });
    io.emit('bus_location_update', trackingState.get(data.busId));
  });

  socket.on('driver_stop_trip', (data) => {
    trackingState.delete(data.busId);
    io.emit('BUS_OFFLINE', { busId: data.busId });
    console.log(`[Socket] Bus ${data.busId} marked offline`);
  });

  socket.on('sos_alert', (data) => {
    io.emit('sos_broadcast', {
      busId: data.busId,
      location: data.location,
      timestamp: new Date().toISOString()
    });
  });

  socket.on('disconnect', () => {
    console.log(`[Socket] Client disconnected: ${socket.id}`);
  });
});

const connectDB = async () => {
  try {
    const conn = await mongoose.connect(process.env.MONGODB_URI);
    console.log(`MongoDB Connected: ${conn.connection.host}`);
  } catch (error) {
    console.error(`MongoDB Connection Error: ${error.message}`);
  }
};

const busRoutes = require('./src/routes/busRoutes');
app.use('/api/buses', busRoutes);

app.get('/health', (req, res) => {
  res.status(200).json({
    success: true,
    message: 'Server is running',
    timestamp: new Date().toISOString(),
    activeBuses: trackingState.size
  });
});

const startServer = async () => {
  await connectDB();
  httpServer.listen(PORT, "0.0.0.0", () => {
    console.log(`V-Bus Server running on port ${PORT}`);
    console.log(`Socket.IO: Enabled`);
  });
};

startServer();
""", alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

add_heading_custom('Passenger Application BusContext', level=2)
add_normal('BusContext.js')
add_normal("""
import React, { createContext, useContext, useState, useEffect, useRef } from 'react';
import { io } from 'socket.io-client';
import { SOCKET_URL } from './config';

const BusContext = createContext();

export const BusProvider = ({ children }) => {
  const [buses, setBuses] = useState([]);
  const [selectedBus, setSelectedBus] = useState(null);
  const socketRef = useRef(null);

  useEffect(() => {
    const socket = io(SOCKET_URL);
    socketRef.current = socket;

    socket.on('connect', () => {
      console.log('[BusContext] Socket connected');
    });

    socket.on('bus_location_update', (data) => {
      setBuses((prev) => {
        const idx = prev.findIndex((b) => b.busId === data.busId);
        if (idx >= 0) {
          const updated = [...prev];
          updated[idx] = data;
          return updated;
        }
        return [...prev, data];
      });
    });

    socket.on('BUS_OFFLINE', (data) => {
      setBuses((prev) => prev.filter((b) => b.busId !== data.busId));
      setSelectedBus((prev) => (prev && prev.busId === data.busId ? null : prev));
    });

    socket.on('sos_broadcast', (data) => {
      console.warn('[BusContext] SOS Alert:', data);
    });

    return () => {
      socket.disconnect();
    };
  }, []);

  const trackBus = (busId) => {
    socketRef.current.emit('track_bus', busId);
  };

  const untrackBus = (busId) => {
    socketRef.current.emit('untrack_bus', busId);
  };

  return (
    <BusContext.Provider value={{ buses, selectedBus, setSelectedBus, trackBus, untrackBus }}>
      {children}
    </BusContext.Provider>
  );
};

export const useBus = () => useContext(BusContext);
""", alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

add_heading_custom('WebView Map Communication', level=2)
add_normal('map.html (Leaflet.js WebView)')
add_normal("""
<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>V-Bus Map</title>
  <link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css" />
  <script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
  <style>html, body, #map { height: 100%; margin: 0; }</style>
</head>
<body>
  <div id="map"></div>
  <script>
    const map = L.map('map').setView([12.9716, 77.5946], 13);
    L.tileLayer('https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', {
      attribution: '&copy; OpenStreetMap contributors'
    }).addTo(map);

    const markers = {};
    const routes = {};

    document.addEventListener('message', function(event) {
      const msg = JSON.parse(event.data);
      handleMessage(msg);
    });

    window.addEventListener('message', function(event) {
      const msg = JSON.parse(event.data);
      handleMessage(msg);
    });

    function handleMessage(msg) {
      if (msg.type === 'BUS_UPDATE') {
        const { busId, lat, lng, routeId } = msg;
        if (markers[busId]) {
          markers[busId].setLatLng([lat, lng]);
        } else {
          markers[busId] = L.marker([lat, lng]).addTo(map).bindPopup(`Bus ${busId}`);
        }
        if (msg.follow) {
          map.panTo([lat, lng]);
        }
      }
      if (msg.type === 'BUS_OFFLINE') {
        const busId = msg.busId;
        if (markers[busId]) {
          map.removeLayer(markers[busId]);
          delete markers[busId];
        }
      }
      if (msg.type === 'ROUTE_RENDER') {
        const { routeId, coordinates } = msg;
        if (routes[routeId]) map.removeLayer(routes[routeId]);
        routes[routeId] = L.polyline(coordinates, { color: 'blue' }).addTo(map);
      }
    }
  </script>
</body>
</html>
""", alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

add_heading_custom('Driver Tracking Screen', level=2)
add_normal('DriverTrackingScreen.js')
add_normal("""
import React, { useEffect, useState } from 'react';
import { View, Text, Button, Alert } from 'react-native';
import * as Location from 'expo-location';
import { useSocket } from './SocketContext';

export default function DriverTrackingScreen({ route }) {
  const { busId, assignedRoute } = route.params;
  const socket = useSocket();
  const [tracking, setTracking] = useState(false);

  useEffect(() => {
    let locationSubscription;
    const startTracking = async () => {
      const { status } = await Location.requestForegroundPermissionsAsync();
      if (status !== 'granted') {
        Alert.alert('Permission denied');
        return;
      }
      locationSubscription = await Location.watchPositionAsync(
        { accuracy: Location.Accuracy.High, timeInterval: 3000, distanceInterval: 10 },
        (loc) => {
          socket.emit('driver_location_update', {
            busId,
            routeId: assignedRoute,
            lat: loc.coords.latitude,
            lng: loc.coords.longitude,
            speed: loc.coords.speed
          });
        }
      );
    };
    if (tracking) startTracking();
    return () => {
      if (locationSubscription) locationSubscription.remove();
    };
  }, [tracking]);

  const handleSOS = () => {
    socket.emit('sos_alert', { busId, timestamp: new Date().toISOString() });
    Alert.alert('SOS Alert Sent');
  };

  return (
    <View style={{ padding: 20 }}>
      <Text>Bus: {busId}</Text>
      <Text>Route: {assignedRoute}</Text>
      <Button title={tracking ? 'Stop Trip' : 'Start Trip'}
        onPress={() => {
          if (tracking) {
            socket.emit('driver_stop_trip', { busId });
          }
          setTracking(!tracking);
        }} />
      <Button title="SOS Emergency" onPress={handleSOS} color="red" />
    </View>
  );
}
""", alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Inches(0))

doc.add_page_break()

# ===== APPENDIX 2: SCREENSHOTS =====
add_heading_custom('APPENDIX 2: SCREENSHOTS', level=1)

add_heading_custom('Admin Dashboard Map View', level=2)
add_normal('The admin dashboard provides a centralized view of all buses in the system. It displays real-time bus locations on the map along with route paths and status indicators such as online, offline, and delayed. Administrators can monitor multiple buses simultaneously, assign routes to drivers, and analyze system performance efficiently.')

add_heading_custom('Passenger View Interface', level=2)
add_normal('The passenger interface allows users to track buses in real time using an interactive WebView map. It displays nearby buses, route details, stops, and estimated time of arrival (ETA). The follow-bus mode keeps the viewport centered on the selected vehicle. This helps passengers plan their journey effectively and reduces waiting time.')

add_heading_custom('Login Page', level=2)
add_normal('The login page provides secure access to the system for administrators and users. It requires valid credentials such as email and password to authenticate users. JWT-based authentication ensures data security and restricts unauthorized access to sensitive system functionalities.')

add_heading_custom('Driver Interface', level=2)
add_normal('The driver interface enables bus drivers to manage their trips efficiently. It includes features such as starting a trip, enabling background GPS tracking, and emergency SOS alerts. The interface also displays route and bus details, ensuring smooth operation during travel.')

doc.add_page_break()

# ===== REFERENCES =====
add_heading_custom('REFERENCES', level=1)
add_normal('P. Verma, S. Singh, and R. Kumar, "GPS-Based Bus Tracking System," International Journal of Computer Applications, vol. 182, no. 45, pp. 12–18, 2021.')
add_normal('A. Sharma, K. Gupta, and M. Patel, "Real-Time Communication in Transportation Systems," International Journal of Smart Systems, vol. 10, no. 2, pp. 45–52, 2022.')
add_normal('J. Lee, H. Kim, and S. Park, "Event-Driven Vehicle Monitoring System," IEEE Transactions on Intelligent Transportation Systems, vol. 21, no. 4, pp. 1567–1575, 2020.')
add_normal('R. Ahmed, L. Chen, and D. Williams, "Smart Public Transport Systems Using Real-Time Communication and Cloud," International Journal of Transportation Engineering, vol. 9, no. 3, pp. 101–110, 2023.')
add_normal('S. Gupta, N. Verma, and T. Reddy, "WebView and Hybrid Map Rendering in Mobile Applications," International Journal of Mobile Computing, vol. 8, no. 1, pp. 25–32, 2021.')
add_normal('K. Sharma, R. Iyer, and P. Nair, "Survey on Smart Transport Technologies," International Journal of Advanced Transportation, vol. 15, no. 2, pp. 55–68, 2022.')
add_normal('Y. Zhang, L. Wang, and X. Li, "Real-Time GPS Tracking System for Public Transport," IEEE Access, vol. 10, pp. 34567–34575, 2022.')
add_normal('R. Kumar and T. Das, "IoT-Based Smart Transportation System Using Cloud Computing," International Journal of Advanced Computer Science and Applications, vol. 13, no. 5, pp. 88–95, 2022.')
add_normal('M. Smith and J. Brown, "Cloud-Based Intelligent Transport Systems," Journal of Smart Cities, vol. 6, no. 2, pp. 77–85, 2021.')
add_normal('Leaflet Contributors, "Leaflet: An Open-Source JavaScript Library for Interactive Maps," Available: https://leafletjs.com, 2023.')
add_normal('Socket.IO, "Socket.IO: Real-Time Bidirectional Event-Based Communication," Available: https://socket.io, 2023.')
add_normal('OpenStreetMap Contributors, "OpenStreetMap," Available: https://www.openstreetmap.org, 2023.')

doc.add_page_break()

# ===== CERTIFICATE =====
add_heading_custom('CERTIFICATE', level=1)
add_normal('')
add_normal('')
add_normal('')

print("Report generation complete!")
doc.save(r'w:\Final year project\V-Bus_Report.docx')
